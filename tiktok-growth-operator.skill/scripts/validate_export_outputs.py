from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries
from render_scene_report import localized_header_text
from text_normalization import read_json_file, write_json_file
from validator_runtime import create_validator_runtime


VISIBLE_TEXT_MOJIBAKE_PATTERNS = [
    "\ufffd",
    "â€™",
    "â€˜",
    "â€œ",
    "â€\x9d",
    "â€“",
    "â€”",
    "â€¦",
    "Â ",
]

VISIBLE_ABSOLUTE_PATH_RE = re.compile(r"(?<![A-Za-z])[A-Za-z]:[\\/][^\s|]+")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render representative scene reports and validate Markdown, DOCX, and XLSX export structure."
    )
    parser.add_argument(
        "--output-root",
        default="",
        help="Optional explicit output root for generated validation artifacts.",
    )
    return parser.parse_args()


def skill_root() -> Path:
    return Path(__file__).resolve().parents[1]


def default_output_root() -> Path:
    return create_validator_runtime(skill_root(), "export")


def validation_fixture_root() -> Path:
    return skill_root() / "testdata" / "validation"


def resolve_fixture_path(*candidates: Path) -> Path:
    for candidate in candidates:
        if candidate.exists():
            return candidate
    searched = [str(candidate) for candidate in candidates]
    raise FileNotFoundError(f"No export-validation fixture found. Checked: {searched}")


def run_render(input_json: Path, output_dir: Path) -> dict:
    command = [
        sys.executable,
        str(skill_root() / "scripts" / "render_scene_report.py"),
        "--input",
        str(input_json),
        "--output-dir",
        str(output_dir),
        "--formats",
        "md,docx,xlsx",
    ]
    completed = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", check=True)
    return json.loads(completed.stdout)


def find_visible_text_mojibake(lines: list[str]) -> list[str]:
    findings: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if any(pattern in stripped for pattern in VISIBLE_TEXT_MOJIBAKE_PATTERNS):
            findings.append(stripped)
    return findings


def find_visible_absolute_paths(lines: list[str]) -> list[str]:
    findings: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        match = VISIBLE_ABSOLUTE_PATH_RE.search(stripped)
        if match:
            findings.append(match.group(0))
    return findings


def validate_markdown(path: Path, report: dict) -> dict:
    text = path.read_text(encoding="utf-8-sig")
    visible_lines = text.splitlines()
    mojibake_findings = find_visible_text_mojibake(visible_lines)
    if mojibake_findings:
        raise AssertionError(f"Markdown contains mojibake-like visible text in {path.name}: {mojibake_findings[0]}")
    absolute_path_findings = find_visible_absolute_paths(visible_lines)
    if absolute_path_findings:
        raise AssertionError(
            f"Markdown contains visible absolute local paths in {path.name}: {absolute_path_findings[0]}"
        )
    required_string_groups = [
        ("## Working Context", "## 任务上下文"),
        ("## Executive Summary", "## 执行摘要"),
    ]
    missing = ["/".join(group) for group in required_string_groups if not any(value in text for value in group)]
    if missing:
        raise AssertionError(f"Markdown missing expected sections in {path.name}: {', '.join(missing)}")

    has_execution = payload_has_execution_template(report)
    if has_execution:
        execution_required_groups = [
            ("## Direct-Use Template", "## 直接执行模板"),
            ("### Variable Inputs", "### 可变输入"),
            ("### Workflow Steps", "### 执行步骤"),
            ("### Output Checklist", "### 交付检查清单"),
        ]
        execution_missing = ["/".join(group) for group in execution_required_groups if not any(value in text for value in group)]
        if execution_missing:
            raise AssertionError(
                f"Markdown missing execution-template sections in {path.name}: {', '.join(execution_missing)}"
            )
    return {
        "visible_text_mojibake": False,
        "visible_absolute_paths": False,
        "has_required_sections": True,
        "has_execution_template": has_execution,
    }


def payload_has_execution_template(report: dict) -> bool:
    execution_template = report.get("execution_template", {}) or {}
    return any(
        execution_template.get(key)
        for key in [
            "recommended_request",
            "recommended_request_zh",
            "recommended_runner_args",
            "variable_inputs",
            "codex_prompt_scaffold",
            "codex_prompt_scaffold_zh",
            "workflow_steps",
            "output_checklist",
        ]
    )


def workbook_sheet_alias(workbook, *candidates: str):
    for name in candidates:
        if name in workbook.sheetnames:
            return workbook[name], name
    raise AssertionError(f"Missing sheets: {', '.join(candidates)}")


def validate_workbook(path: Path, report: dict) -> dict:
    workbook = load_workbook(path)
    summary, summary_name = workbook_sheet_alias(workbook, "Summary", "总览")
    overview, overview_name = workbook_sheet_alias(workbook, "Section Overview", "章节概览")
    section_index, section_index_name = workbook_sheet_alias(workbook, "Section Index", "章节导航")
    if overview["A4"].hyperlink is None:
        raise AssertionError(f"Section Overview first section link missing in {path.name}")
    if section_index["B4"].hyperlink is None or section_index["C4"].hyperlink is None:
        raise AssertionError(f"Section Index navigation link missing in {path.name}")

    target_sheet = section_index["C4"].value
    if target_sheet not in workbook.sheetnames:
        raise AssertionError(f"Section Index target sheet missing in {path.name}: {target_sheet}")
    back_link = workbook[target_sheet]["A2"].hyperlink
    if back_link is None or back_link.target not in {f"#'{section_index_name}'!A1", "#'Section Index'!A1", "#'章节导航'!A1"}:
        raise AssertionError(f"Section sheet back-link invalid in {path.name}")

    has_execution = payload_has_execution_template(report)
    if has_execution:
        execution_sheet, execution_name = workbook_sheet_alias(workbook, "Execution Template", "执行模板")
        summary_template_link = None
        for row in range(7, summary.max_row + 1):
            if summary.cell(row=row, column=1).value in {"Execution Template", "执行模板"}:
                summary_template_link = summary.cell(row=row, column=2).hyperlink
                break
        if summary_template_link is None or summary_template_link.target not in {f"#'{execution_name}'!A1", "#'Execution Template'!A1", "#'执行模板'!A1"}:
            raise AssertionError(f"Summary execution-template link missing in {path.name}")
        if execution_sheet["A1"].value not in {"Direct-Use Template", "执行模板"}:
            raise AssertionError(f"Execution Template title missing in {path.name}")
        if execution_sheet["A2"].hyperlink is None or execution_sheet["A2"].hyperlink.target not in {f"#'{summary_name}'!A1", "#'Summary'!A1", "#'总览'!A1"}:
            raise AssertionError(f"Execution Template back-link missing in {path.name}")

    if not summary.tables:
        raise AssertionError(f"Summary table missing in {path.name}")
    if not overview.tables:
        raise AssertionError(f"Section Overview table missing in {path.name}")
    if not section_index.tables:
        raise AssertionError(f"Section Index table missing in {path.name}")
    for cell in ["A3", "B3", "C3", "D3", "E3", "F3"]:
        if not summary[cell].value:
            raise AssertionError(f"Summary metric card missing at {cell} in {path.name}")
    for cell in ["A4", "B4", "C4", "D4", "E4", "F4"]:
        if not summary[cell].value:
            raise AssertionError(f"Summary quality card missing at {cell} in {path.name}")

    section_sheet_names = [
        section_index.cell(row=row, column=3).value
        for row in range(4, section_index.max_row + 1)
        if section_index.cell(row=row, column=3).value
    ]
    if len(section_sheet_names) != len(report["sections"]):
        raise AssertionError(
            f"Section Index count mismatch in {path.name}: expected {len(report['sections'])}, got {len(section_sheet_names)}"
        )
    for section, sheet_name in zip(report["sections"], section_sheet_names):
        if sheet_name not in workbook.sheetnames:
            raise AssertionError(f"Section sheet missing in {path.name}: {sheet_name}")
        if not section["table"]["headers"]:
            continue
        section_ws = workbook[sheet_name]
        if not section_ws.tables:
            raise AssertionError(f"Section table missing in {path.name}: {sheet_name}")
        first_table = next(iter(section_ws.tables.values()))
        min_col, min_row, max_col, _ = range_boundaries(first_table.ref)
        expected_headers = section["table"]["headers"]
        localized_expected_headers = [localized_header_text(header) for header in expected_headers]
        actual_headers = [
            section_ws.cell(row=min_row, column=column).value or ""
            for column in range(min_col, max_col + 1)
        ]
        if (max_col - min_col + 1) != len(expected_headers):
            raise AssertionError(
                f"Section table width mismatch in {path.name}: {sheet_name} expected {len(expected_headers)} columns, got {max_col - min_col + 1}"
            )
        if actual_headers != expected_headers and actual_headers != localized_expected_headers:
            raise AssertionError(
                f"Section table headers mismatch in {path.name}: {sheet_name} expected {expected_headers} or {localized_expected_headers}, got {actual_headers}"
            )

    visible_lines: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for value in row:
                if isinstance(value, str):
                    visible_lines.append(value)
    mojibake_findings = find_visible_text_mojibake(visible_lines)
    if mojibake_findings:
        raise AssertionError(
            f"Workbook contains mojibake-like visible text in {path.name}: {mojibake_findings[0]}"
        )
    absolute_path_findings = find_visible_absolute_paths(visible_lines)
    if absolute_path_findings:
        raise AssertionError(
            f"Workbook contains visible absolute local paths in {path.name}: {absolute_path_findings[0]}"
        )

    return {
        "sheets": workbook.sheetnames,
        "summary_tables": list(summary.tables.keys()),
        "overview_tables": list(overview.tables.keys()),
        "index_tables": list(section_index.tables.keys()),
        "first_section_sheet": target_sheet,
        "summary_quality_values": [summary[cell].value for cell in ["A4", "B4", "C4", "D4", "E4", "F4"]],
        "section_index_sheet_names": section_sheet_names,
        "has_execution_template_sheet": has_execution and any(name in workbook.sheetnames for name in ["Execution Template", "执行模板"]),
        "visible_text_mojibake": False,
        "visible_absolute_paths": False,
    }


def validate_docx(path: Path, report: dict) -> dict:
    with ZipFile(path) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")

    required_string_groups = [
        ("Contents", "目录"),
        ("Section Overview", "章节概览"),
        ("Back to Contents", "返回目录"),
        ("Back to Section Overview", "返回章节概览"),
    ]
    missing_strings = ["/".join(group) for group in required_string_groups if not any(value in document_xml for value in group)]
    if missing_strings:
        raise AssertionError(f"DOCX missing expected navigation text in {path.name}: {', '.join(missing_strings)}")
    if "contents_anchor" not in document_xml or "section_overview" not in document_xml:
        raise AssertionError(f"DOCX missing expected bookmarks in {path.name}")
    has_embedded_media = any(name.startswith("word/media/") for name in names)
    if has_embedded_media and "Figure 1." not in document_xml:
        raise AssertionError(f"DOCX missing expected figure caption text in {path.name}")
    has_execution = payload_has_execution_template(report)
    if has_execution:
        for group in [("Direct-Use Template", "执行模板"), ("Workflow Steps", "执行步骤"), ("Output Checklist", "交付检查清单")]:
            if not any(value in document_xml for value in group):
                raise AssertionError(f"DOCX missing execution-template content in {path.name}: {'/'.join(group)}")

    document = Document(path)
    visible_lines = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                visible_lines.append(cell.text)
    for section in document.sections:
        visible_lines.extend(paragraph.text for paragraph in section.header.paragraphs)
        visible_lines.extend(paragraph.text for paragraph in section.footer.paragraphs)
    mojibake_findings = find_visible_text_mojibake(visible_lines)
    if mojibake_findings:
        raise AssertionError(f"DOCX contains mojibake-like visible text in {path.name}: {mojibake_findings[0]}")
    absolute_path_findings = find_visible_absolute_paths(visible_lines)
    if absolute_path_findings:
        raise AssertionError(f"DOCX contains visible absolute local paths in {path.name}: {absolute_path_findings[0]}")

    return {
        "has_contents": True,
        "has_section_overview": True,
        "has_back_links": True,
        "has_figure_caption": (not has_embedded_media) or ("Figure 1." in document_xml),
        "has_embedded_media": has_embedded_media,
        "has_bookmarks": True,
        "has_execution_template": (not has_execution) or any(value in document_xml for value in ["Direct-Use Template", "执行模板"]),
        "visible_text_mojibake": False,
        "visible_absolute_paths": False,
    }


def synthetic_duplicate_heading_report() -> dict:
    return {
        "metadata": {
            "scene": "17",
            "project": "Synthetic Duplicate Heading Check",
            "title": "Synthetic Duplicate Heading Check",
            "status": "draft",
            "generated_at": "2026-05-04",
            "scene_title": "Synthetic Regression",
            "deliverable_type": "validation",
            "scenario_file": "synthetic.json",
            "scene_slug": "synthetic-regression",
        },
        "working_context": {
            "summary": "Synthetic report for duplicate section heading validation.",
            "inputs": ["fixture-a"],
            "minimum_evidence": [],
            "ideal_evidence": [],
            "constraints": [],
            "requested_outputs": ["docx", "xlsx"],
            "ready_checklist": ["ok"],
        },
        "executive_summary": {
            "conclusion": "Synthetic regression passed.",
            "why_it_matters": "Stable sheet naming must survive duplicate headings.",
            "next_action": "Keep exporter deterministic.",
            "confidence": "high",
        },
        "operator_guide": {
            "operator_checklist": ["Render and inspect workbook links."],
            "common_failure_modes": ["Broken overview links when titles collide."],
        },
        "sections": [
            {
                "heading": "Repeated Section",
                "instruction": "First repeated heading.",
                "paragraphs": ["Alpha"],
                "bullets": [],
                "numbered": [],
                "table": {"title": "", "headers": [], "rows": []},
            },
            {
                "heading": "Repeated Section",
                "instruction": "Second repeated heading.",
                "paragraphs": ["Beta"],
                "bullets": [],
                "numbered": [],
                "table": {"title": "", "headers": [], "rows": []},
            },
        ],
        "evidence": [],
        "assets": [{"label": "Fixture Image", "path": "", "note": "Synthetic asset row for caption validation."}],
        "notes": [],
        "sources": [],
    }


def synthetic_sparse_report() -> dict:
    return {
        "metadata": {
            "scene": "15",
            "project": "Synthetic Sparse Section Check",
            "title": "Synthetic Sparse Section Check",
            "status": "draft",
            "generated_at": "2026-05-04",
            "scene_title": "Synthetic Regression",
            "deliverable_type": "validation",
            "scenario_file": "synthetic.json",
            "scene_slug": "synthetic-regression",
        },
        "working_context": {
            "summary": "Synthetic sparse report for empty-section validation.",
            "inputs": [],
            "minimum_evidence": [],
            "ideal_evidence": [],
            "constraints": [],
            "requested_outputs": [],
            "ready_checklist": [],
        },
        "executive_summary": {
            "conclusion": "Sparse render stays stable.",
            "why_it_matters": "Empty sections should not break tables or links.",
            "next_action": "Keep placeholder handling intact.",
            "confidence": "medium",
        },
        "operator_guide": {
            "operator_checklist": [],
            "common_failure_modes": [],
        },
        "sections": [
            {
                "heading": "Empty Section",
                "instruction": "",
                "paragraphs": [],
                "bullets": [],
                "numbered": [],
                "table": {"title": "", "headers": [], "rows": []},
            }
        ],
        "evidence": [],
        "assets": [{"label": "Sparse Fixture Image", "path": "", "note": "Synthetic asset row for caption validation."}],
        "notes": [],
        "sources": [],
    }


def synthetic_execution_template_report() -> dict:
    return {
        "metadata": {
            "scene": "01",
            "project": "Synthetic Execution Template Check",
            "title": "Synthetic Execution Template Check",
            "status": "draft",
            "generated_at": "2026-05-07",
            "scene_title": "Synthetic Regression",
            "deliverable_type": "validation",
            "scenario_file": "synthetic.json",
            "scene_slug": "synthetic-regression",
        },
        "working_context": {
            "summary": "Synthetic report for execution-template export validation.",
            "inputs": ["fixture-template"],
            "minimum_evidence": ["request"],
            "ideal_evidence": [],
            "constraints": [],
            "requested_outputs": ["docx", "xlsx"],
            "ready_checklist": ["template present"],
        },
        "executive_summary": {
            "conclusion": "Execution template exports are present.",
            "why_it_matters": "Reusable prompt/workflow instructions must survive DOCX/XLSX rendering.",
            "next_action": "Keep parity between markdown and rich exports.",
            "confidence": "high",
        },
        "operator_guide": {
            "operator_checklist": ["Verify request scaffold survives export."],
            "common_failure_modes": ["Markdown-only template instructions."],
        },
        "execution_template": {
            "recommended_request": "Run Scene 01 against a real TikTok profile and return a shortlist.",
            "recommended_request_zh": "运行 Scene 01，针对真实 TikTok 账号输出候选短名单。",
            "recommended_runner_args": ["--scene 01", "--qualified-count 5"],
            "variable_inputs": [
                {"name": "profile_url", "meaning": "TikTok profile to analyze", "example": "https://www.tiktok.com/@mustsharenews", "required": "yes"}
            ],
            "codex_prompt_scaffold": ["Use real TikTok evidence only.", "Return shortlist plus why-selected rationale."],
            "codex_prompt_scaffold_zh": ["只使用真实 TikTok 证据。", "输出 shortlist 和 why-selected 结论。"],
            "workflow_steps": ["Collect profile posts.", "Rank qualified candidates.", "Write shortlist report."],
            "output_checklist": ["Shortlist exists.", "Evidence links are attached."],
        },
        "sections": [
            {
                "heading": "Template Validation",
                "instruction": "Confirm the execution template is exported.",
                "paragraphs": ["This synthetic fixture exists only to validate export parity."],
                "bullets": [],
                "numbered": [],
                "table": {"title": "", "headers": [], "rows": []},
            }
        ],
        "evidence": [],
        "assets": [],
        "notes": [],
        "sources": [],
    }


def synthetic_wide_table_report() -> dict:
    return {
        "metadata": {
            "scene": "03",
            "project": "Synthetic Wide Table Check",
            "title": "Synthetic Wide Table Check",
            "status": "draft",
            "generated_at": "2026-05-08",
            "scene_title": "Synthetic Regression",
            "deliverable_type": "validation",
            "scenario_file": "synthetic.json",
            "scene_slug": "synthetic-regression",
        },
        "working_context": {
            "summary": "Validate wide-table export and path cleanup from D:\\synthetic\\captures\\wide-table\\summary.json.",
            "inputs": [
                "Capture root: D:\\synthetic\\captures\\wide-table",
                "Market: US",
            ],
            "minimum_evidence": ["Wide section table fixture"],
            "ideal_evidence": [],
            "constraints": [],
            "requested_outputs": ["docx", "xlsx"],
            "ready_checklist": ["Wide table included"],
        },
        "executive_summary": {
            "conclusion": "Wide section tables should keep all columns.",
            "why_it_matters": "Scene exports now rely on 5-8 column tables for platform-style deliverables.",
            "next_action": "Keep section-sheet widths and headers stable.",
            "confidence": "high",
        },
        "operator_guide": {
            "operator_checklist": ["Verify every section-table header survives the rich export."],
            "common_failure_modes": ["Section-sheet merge logic collapsing wide tables back to four columns."],
        },
        "sections": [
            {
                "heading": "Wide Table Validation",
                "instruction": "Keep all table headers visible and preserve cleaned path text.",
                "paragraphs": [
                    "This fixture exists to validate section-sheet widths, DOCX table widths, and visible path cleanup."
                ],
                "bullets": [
                    "Visible text must not leak raw local absolute paths.",
                    "The section sheet must preserve six table columns.",
                ],
                "numbered": [],
                "table": {
                    "title": "Wide Validation Grid",
                    "headers": [
                        "Signal",
                        "Source Product",
                        "Hook",
                        "Core Topic",
                        "Commerce Angle",
                        "Next Action",
                    ],
                    "rows": [
                        [
                            "High replay rate",
                            "Orange Cat Plush",
                            "Instant visual gag",
                            "Pet comfort humor",
                            "Giftable impulse",
                            "Route to Scene 03 shortlist",
                        ],
                        [
                            "Comment spike",
                            "Desk toy",
                            "Unexpected motion demo",
                            "Stress relief",
                            "Problem-solution proof",
                            "Write follow-up breakdown brief",
                        ],
                    ],
                },
            }
        ],
        "evidence": [
            {
                "label": "Synthetic Capture",
                "detail": "Fixture derived from D:\\synthetic\\captures\\wide-table\\ranked_videos.json.",
                "source": "synthetic-wide-table",
            }
        ],
        "assets": [],
        "notes": ["Path cleanup should compact local roots into display-safe relative paths."],
        "sources": ["synthetic://wide-table-check"],
    }


def run_synthetic_render(report: dict, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", suffix=".json", delete=False, encoding="utf-8") as handle:
        json.dump(report, handle, ensure_ascii=False, indent=2)
        temp_path = Path(handle.name)
    try:
        return run_render(temp_path, output_dir)
    finally:
        temp_path.unlink(missing_ok=True)


def main() -> None:
    args = parse_args()
    root = Path(args.output_root) if args.output_root else default_output_root()
    root.mkdir(parents=True, exist_ok=True)

    fixtures = [
        {
            "name": "scene15",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-15-validation-scene15-capture.json",
                skill_root() / "tmp" / "20260504_validation_capture_scene15" / "scene-15" / "scene-15-validation-scene15-capture.json",
                skill_root() / "tmp" / "20260504_capture_runner_scene15" / "scene-15" / "scene-15-scene15-capture-run.json",
            ),
        },
        {
            "name": "scene17",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-17-validation-routed-capture.json",
                skill_root() / "tmp" / "20260504_validation_routed_capture" / "scene-17" / "scene-17-validation-routed-capture.json",
                skill_root() / "tmp" / "20260504_capture_runner_scene17" / "scene-17" / "scene-17-tiktok-official-capture-run.json",
            ),
        },
        {
            "name": "scene02_real",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-02-validation-scene02-patrol.json",
                skill_root() / "tmp" / "20260507_validation_capture_scene02" / "scene02-run" / "scene-02" / "scene-02-validation-scene02-patrol.json",
            ),
        },
        {
            "name": "scene03_real",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-03-patrol-scene03-handoff.json",
                skill_root() / "tmp" / "20260507_validation_capture_scene02" / "scene02-run" / "scene-03-from-patrol" / "scene-03" / "scene-03-patrol-scene03-handoff.json",
            ),
        },
        {
            "name": "scene08_real",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-08-validation-auto-capture.json",
                skill_root() / "tmp" / "20260504_validation_capture_scene_auto" / "scene-08" / "scene-08-auto-validation-auto-capture.json",
            ),
        },
        {
            "name": "scene18_real",
            "input": resolve_fixture_path(
                validation_fixture_root() / "reports" / "scene-18-spotcheck-scene18-rich.json",
                Path(".codex-tmp") / "spotcheck-scene18-rich" / "operator-run" / "scene-18" / "scene-18-spotcheck-scene18-rich.json",
            ),
        },
    ]
    synthetic_fixtures = [
        {"name": "synthetic_duplicate_heading", "report": synthetic_duplicate_heading_report()},
        {"name": "synthetic_sparse_section", "report": synthetic_sparse_report()},
        {"name": "synthetic_execution_template", "report": synthetic_execution_template_report()},
        {"name": "synthetic_wide_table", "report": synthetic_wide_table_report()},
    ]

    results = []
    for fixture in fixtures:
        output_dir = root / fixture["name"]
        rendered = run_render(fixture["input"], output_dir)
        report = read_json_file(fixture["input"])
        markdown_check = validate_markdown(Path(rendered["md"]), report)
        workbook_check = validate_workbook(Path(rendered["xlsx"]), report)
        docx_check = validate_docx(Path(rendered["docx"]), report)
        results.append(
            {
                "name": fixture["name"],
                "input": str(fixture["input"]),
                "output_dir": str(output_dir),
                "rendered": rendered,
                "markdown_check": markdown_check,
                "workbook_check": workbook_check,
                "docx_check": docx_check,
            }
        )

    for fixture in synthetic_fixtures:
        output_dir = root / fixture["name"]
        rendered = run_synthetic_render(fixture["report"], output_dir)
        markdown_check = validate_markdown(Path(rendered["md"]), fixture["report"])
        workbook_check = validate_workbook(Path(rendered["xlsx"]), fixture["report"])
        docx_check = validate_docx(Path(rendered["docx"]), fixture["report"])
        if fixture["name"] == "synthetic_duplicate_heading":
            sheet_names = workbook_check["section_index_sheet_names"]
            if len(sheet_names) != len(set(sheet_names)):
                raise AssertionError("Duplicate-heading synthetic fixture produced non-unique section sheet names")
        if fixture["name"] == "synthetic_sparse_section":
            quality_values = [str(value) for value in workbook_check["summary_quality_values"]]
            if "1\nBlank Sections" not in quality_values and "1\n空白章节" not in quality_values:
                raise AssertionError("Sparse synthetic fixture did not flag blank sections in Summary quality cards")
        if fixture["name"] == "synthetic_execution_template":
            if not markdown_check["has_execution_template"]:
                raise AssertionError("Execution-template synthetic fixture did not preserve markdown template content")
            if not workbook_check["has_execution_template_sheet"]:
                raise AssertionError("Execution-template synthetic fixture did not produce the Execution Template sheet")
            if not docx_check["has_execution_template"]:
                raise AssertionError("Execution-template synthetic fixture did not produce DOCX execution-template content")
        results.append(
            {
                "name": fixture["name"],
                "input": "synthetic",
                "output_dir": str(output_dir),
                "rendered": rendered,
                "markdown_check": markdown_check,
                "workbook_check": workbook_check,
                "docx_check": docx_check,
            }
        )

    summary_path = root / "validation_summary.json"
    write_json_file(summary_path, results)
    print(json.dumps({"status": "ok", "summary": str(summary_path), "runs": results}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
