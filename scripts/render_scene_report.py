from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import qn as qn_shared
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

from feishu_naming import build_report_title, localize_project_text, scene_label_zh
from generate_scene_report import build_report_payload, load_catalog, render_markdown_from_payload, resolve_scene
from text_normalization import normalize_nested, normalize_text, read_json_file, read_utf8_text, write_utf8_text


TITLE_FILL = PatternFill(fill_type="solid", fgColor="1F4E78")
HEADER_FILL = PatternFill(fill_type="solid", fgColor="D9EAF7")
LABEL_FILL = PatternFill(fill_type="solid", fgColor="EEF4FB")
SUCCESS_FILL = PatternFill(fill_type="solid", fgColor="2F7D32")
WARNING_FILL = PatternFill(fill_type="solid", fgColor="D9A404")
DANGER_FILL = PatternFill(fill_type="solid", fgColor="B42318")
WRAP = Alignment(wrap_text=True, vertical="top")
THIN_SIDE = Side(style="thin", color="B8CBE0")
THIN_BORDER = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)
HYPERLINK_FONT = Font(color="0563C1", underline="single")
SHEET_TITLES = {
    "summary": "总览",
    "section_overview": "章节概览",
    "section_index": "章节导航",
    "execution_template": "执行模板",
    "context_lists": "上下文清单",
    "evidence": "证据",
    "assets": "资产",
    "notes": "备注",
    "sources": "来源",
    "operator_guide": "操作指引",
}
RESERVED_SHEET_TITLES = set(SHEET_TITLES.values())

STATUS_LABELS_ZH = {
    "draft": "草稿",
    "final": "成品",
    "done": "完成",
    "complete": "完成",
    "confirmed": "已确认",
    "imported": "已导入",
}

CONTEXT_LABELS_ZH = {
    "inputs": "输入材料",
    "minimum_evidence": "最低证据要求",
    "ideal_evidence": "理想证据补充",
    "constraints": "约束条件",
    "requested_outputs": "目标交付",
    "ready_checklist": "开跑前检查",
}

COMMON_HEADER_LABELS_ZH = {
    "Source Type": "来源类型",
    "Source ID": "来源标识",
    "Source URL": "来源链接",
    "Time Range": "时间范围",
    "Excerpt": "摘录",
    "Supports": "支撑结论",
    "Label": "标签",
    "Detail": "详情",
    "Source": "来源",
    "Path": "路径",
    "Note": "备注",
    "Scene Type": "场景类型",
    "Visual Content": "画面内容",
    "Spoken / On-Screen Script": "口播 / 画面文案",
    "Role In Conversion": "转化作用",
    "Asset / Talent Needed": "所需素材 / 人物",
    "Evidence Ref": "证据引用",
    "Mechanism Layer": "机制层",
    "Observed Pattern": "观察到的模式",
    "Why It Works": "为什么有效",
    "Failure Mode If Removed": "移除后风险",
    "Observed": "观察结果",
    "Reusable?": "可复用？",
    "Adaptation Note": "改编说明",
    "Confidence": "置信度",
    "Lens": "视角",
    "Strategic Role": "策略作用",
    "Element": "元素",
    "Beat / Shot": "节拍 / 镜头",
    "What Must Happen": "必须发生什么",
    "Purpose": "目的",
    "Subtitle / VO Beat": "字幕 / 口播节拍",
    "Proof Block": "证明模块",
    "Primary Asset Need": "主要素材需求",
    "Primary Risk": "主要风险",
    "Decision Area": "决策领域",
    "Recommendation": "建议",
    "Why": "原因",
    "Base Value Or Improvement Opportunity?": "基础价值 / 改进机会",
    "Price Band": "价格带",
    "Repeated Driver": "重复驱动因素",
    "Repeated Complaint": "重复抱怨",
    "Implication": "启发",
    "Pattern Area": "模式维度",
    "What Repeats": "重复点",
    "Example Evidence": "示例证据",
    "Transferable Pattern": "可迁移模式",
    "Why It Transfers": "为什么可迁移",
    "How To Adapt": "如何改写",
    "Creator-Specific Advantage": "创作者专属优势",
    "Why It Does Not Transfer Cleanly": "为什么不能直接迁移",
    "Lane": "处理通道",
    "What To Do Now": "当前动作",
    "Why This Lane Exists": "设这个通道的原因",
    "Layer": "层级",
    "Source Product": "来源商品",
    "Volume": "数量",
    "Primary Purchase Trigger": "主要购买触发因素",
    "Primary Complaint": "主要抱怨",
    "Cluster Type": "聚类类型",
    "Repeated Phrase / Theme": "重复短语 / 主题",
    "What It Suggests": "说明了什么",
    "Product / Content Implication": "产品 / 内容启发",
    "Why Capture It": "为什么要采集",
    "Required Next Time?": "下次必须补吗？",
    "Why It Matters": "为什么重要",
    "Dimension": "维度",
    "Observed Evidence": "观察证据",
    "Likely Intent": "可能意图",
    "Generator Handoff": "生成器交接字段",
    "Creative Layer": "创意层",
    "What Evidence Supports It": "证据支撑",
    "Asset Dependency": "素材依赖",
    "Block": "模块",
    "Prompt / Brief Content": "提示词 / 制作简报内容",
    "Generator Handoff Field": "生成器交接字段",
    "Shot": "镜头段",
    "Duration": "时长",
    "Scene / Subject": "场景 / 主体",
    "Action": "动作",
    "Voiceover / Overlay": "口播 / 叠字",
    "Asset Need": "素材需求",
    "Adaptation Layer": "改写层",
    "Keep From Reference": "保留参考项",
    "Rewrite For Product": "按产品重写",
    "Asset / Talent Dependency": "素材 / 人物依赖",
    "Open Risk": "未解决风险",
    "Delivery Block": "交付模块",
    "What Must Be Finalized": "必须定稿的内容",
    "Who Uses It": "下游使用方",
    "Blocking Gap": "阻塞缺口",
    "Next Owner": "下一责任人",
}

TABLE_TITLE_LABELS_ZH = {
    "Top Candidate Board": "候选主看板",
    "Scene 03 Handoff Shortlist": "Scene 03 交接短名单",
    "Patrol Table Schema": "巡检表结构",
    "Reusable Daily Summary Template": "复用日报模板",
    "Shortlist": "短名单",
    "Per-Video Breakdown Grid": "逐条视频拆解表",
    "Creation Rules": "创作规则",
    "Timeline Breakdown": "时间轴拆解",
    "Mechanism Breakdown": "机制拆解",
    "Viral Interpretation": "爆点解读",
    "Replication Shot Order": "复刻镜头顺序",
    "Generator Handoff": "生成器交接",
    "Production-Spec Handoff": "制作规格交接",
    "Product-Adapted Brief": "产品适配制作简报",
    "Generator-Ready Brief": "可生成制作简报",
    "Creator Formula Board": "创作者公式面板",
    "Production Handoff": "制作交接",
    "Execution Handoff": "执行交接",
    "Source Product Summary": "来源商品概览",
    "Comment Signal Clusters": "评论信号聚类",
    "Inferred Original Brief Schema": "反推原始制作简报结构",
    "Shot-Level Breakdown": "镜头级拆解",
}

SECTION_HEADING_LABELS_ZH = {
    "Executive Conclusion": "执行结论",
    "High-Level Judgment": "高层判断",
    "Evidence Clusters": "证据聚类",
    "Recommended Action": "建议动作",
    "Open Questions": "待确认问题",
    "Objects To Track": "监控对象",
    "Why They Matter": "为什么值得关注",
    "Fields To Capture Next Time": "下次补采字段",
    "Next Action": "下一步动作",
    "Structure Logic": "结构逻辑",
    "Core Mechanism": "核心机制",
    "Reusable Formula": "可复用公式",
    "Risks And Adaptation Notes": "风险与适配说明",
    "Target": "目标对象",
    "Audience": "目标人群",
    "Message": "核心信息",
    "Creative Constraints": "创作约束",
    "Production Handoff": "制作交接",
    "Production-Spec Handoff": "制作规格交接",
    "Execution Handoff": "执行交接",
    "BGM And Sensory Layer": "BGM 与感官层",
    "Visual And Distribution Signature": "视觉与分发特征",
    "Core Invariant": "核心不变量",
    "Structure": "结构设计",
    "Variable Matrix": "变量矩阵",
    "What To Learn": "学习目标",
    "Expected Effect": "预期效果",
    "Fallback Mode": "降级模式",
    "Template Validation": "模板校验",
    "Wide Table Validation": "宽表校验",
    "Empty Section": "空白章节校验",
    "Repeated Section": "重复章节校验",
}

SECTION_SHEET_LABELS_ZH = {
    "Executive Conclusion": "执行结论",
    "High-Level Judgment": "高层判断",
    "Evidence Clusters": "证据聚类",
    "Recommended Action": "建议动作",
    "Open Questions": "待确认问题",
    "Objects To Track": "监控对象",
    "Why They Matter": "关注理由",
    "Fields To Capture Next Time": "下次补采字段",
    "Next Action": "下一步动作",
    "Structure Logic": "结构逻辑",
    "Core Mechanism": "核心机制",
    "Reusable Formula": "复用公式",
    "Risks And Adaptation Notes": "风险与适配",
    "Target": "目标对象",
    "Audience": "目标人群",
    "Message": "核心信息",
    "Creative Constraints": "创作约束",
    "Production Handoff": "制作交接",
    "Production-Spec Handoff": "规格交接",
    "Execution Handoff": "执行交接",
    "BGM And Sensory Layer": "BGM与感官层",
    "Visual And Distribution Signature": "视觉分发表征",
    "Core Invariant": "核心不变量",
    "Structure": "结构设计",
    "Variable Matrix": "变量矩阵",
    "What To Learn": "学习目标",
    "Expected Effect": "预期效果",
    "Fallback Mode": "降级模式",
    "Template Validation": "模板校验",
    "Wide Table Validation": "宽表校验",
    "Empty Section": "空白章节",
    "Repeated Section": "重复章节",
}


def get_sheet_title(key: str) -> str:
    return SHEET_TITLES[key]


def localize_status_text(value: str) -> str:
    normalized = normalize_text(value).lower()
    if not normalized:
        return ""
    return STATUS_LABELS_ZH.get(normalized, value)


def localized_scene_display(metadata: dict) -> str:
    scene_id = normalize_text(metadata.get("scene", ""))
    scene_name = scene_label_zh(scene_id)
    return f"{scene_id} - {scene_name}".strip(" -")


def localized_report_title(metadata: dict) -> str:
    return build_report_title(metadata.get("project", ""), metadata.get("scene", ""), metadata.get("scene_title", ""))


def localized_section_heading(heading: str) -> str:
    normalized = normalize_text(heading)
    return SECTION_HEADING_LABELS_ZH.get(normalized, normalized or "未命名章节")


def localized_sheet_heading(heading: str) -> str:
    normalized = normalize_text(heading)
    return SECTION_SHEET_LABELS_ZH.get(normalized, normalized or "章节")


def localized_header_text(header: str) -> str:
    normalized = normalize_text(header)
    return COMMON_HEADER_LABELS_ZH.get(normalized, normalized)


def localized_table_title(title: str) -> str:
    normalized = normalize_text(title)
    return TABLE_TITLE_LABELS_ZH.get(normalized, normalized)


SHARED_TEMPLATE_TEXT_REPLACEMENTS = {
    "aggregate_summary.json": "汇总文件（aggregate_summary.json）",
    "profile_summary.json": "账号汇总文件（profile_summary.json）",
    "aggregate_ranked_videos.json": "排序视频文件（aggregate_ranked_videos.json）",
    "aggregate_qualified_videos.json": "达标视频文件（aggregate_qualified_videos.json）",
    "summary.json or aggregate_summary.json": "summary.json 或 aggregate_summary.json",
    "profile_summary.json or summary.json": "profile_summary.json 或 summary.json",
    "ranked_videos.json or aggregate_ranked_videos.json": "ranked_videos.json 或 aggregate_ranked_videos.json",
    "aggregate_qualified_videos.json or qualified_video_links.txt": "aggregate_qualified_videos.json 或 qualified_video_links.txt",
    "Human-readable run or campaign name": "便于识别的运行名或项目名",
    "Target market or locale when the scene depends on one market": "当场景依赖单一市场时的目标市场或地区",
    "Links, screenshots, transcripts, exports, OCR text, or copied notes used as source evidence": "作为源证据使用的链接、截图、转写、导出文件、OCR 文本或摘录备注",
    "What the operator wants this scene to produce": "操作者希望该场景产出的结果",
    "One Viral Concealer Breakdown": "口红爆款单视频拆解",
    "single-video mechanism breakdown and adaptation path": "单视频机制拆解与改编路径",
    "One video link or storyboard": "1 条视频链接或一份分镜摘要",
    "One video link or storyboard summary": "1 条视频链接或一份分镜摘要",
    "One video link, transcript notes, and screenshots by beat": "1 条视频链接、转写笔记与按节拍截图",
    "参考视频 frames, transcript snippets, and pacing notes": "参考视频逐帧画面、转写片段与节奏笔记",
    "Transcript or subtitle notes": "转写稿或字幕笔记",
    "Frame notes or screenshots": "逐帧笔记或截图",
    "Optional basic performance context": "可选的基础表现数据",
    "Screenshots by beat": "按节拍截取的截图",
    "Basic performance context": "基础表现上下文",
    "Audio or BGM cue": "音频或 BGM 线索",
    "Timeline breakdown table": "时间轴拆解表",
    "Video-type classification": "视频类型归类",
    "BGM analysis": "BGM 分析",
    "Three-part viral interpretation": "三段式爆点解读",
    "Reusable mechanism": "可复用机制",
    "Adaptation advice": "改编建议",
    "Separate deep logic from surface style.": "把深层机制与表层风格分开看。",
    "If there is no voiceover, reconstruct the logic from subtitles, actions, cuts, and visual proof.": "如果没有口播，就从字幕、动作、剪辑与视觉证明链里重建逻辑。",
    "Reference video": "参考视频",
    "Screenshots or frame summary": "截图或逐帧摘要",
    "Transcript notes": "转写笔记",
    "Optional user product to adapt onto": "可选的用户产品改写目标",
    "Transcript": "转写稿",
    "Frame-by-frame notes": "逐帧笔记",
    "User product for adaptation": "用于改写的用户产品",
    "Scene or talent notes": "场景或人物备注",
    "One video or visual summary": "1 条视频或一份视觉摘要",
    "If evidence is thin, mark the prompt as low-confidence.": "如果证据偏薄，就把提示词标记为低置信度。",
    "Do not invent hidden production details; keep uncertain fields explicit.": "不要臆造隐藏制作细节，不确定字段要明确写出来。",
    "Inferred original brief": "反推原始制作简报",
    "Generator-ready schema": "可生成结构",
    "Shot-by-shot table": "分镜逐条表",
    "Product-adapted brief": "产品适配制作简报",
    "Field-level confidence flags": "字段级置信度标记",
    "Inferred original brief, Generator-ready schema, Shot-by-shot table, Product-adapted brief, Field-level confidence flags": "反推原始制作简报、可生成结构、分镜逐条表、产品适配制作简报、字段级置信度标记",
    "Inferred original brief, Generator-ready schema, Shot-by-shot table, Product-adapted brief, Field-level confidence flags": "反推原始制作简报、可生成结构、分镜逐条表、产品适配制作简报、字段级置信度标记",
    "Creator Brief Reconstruction": "创作者制作简报反推",
    "reverse-engineered prompt or production brief": "反推提示词与制作简报",
    "Reference video frames, transcript snippets, and pacing notes": "参考视频逐帧画面、转写片段与节奏笔记",
    "One creator account or several videos from one creator": "1 个创作者账号或同一创作者的多条视频",
    "Top videos": "高表现视频",
    "Transcripts": "转写稿",
    "Performance notes": "表现备注",
    "Optional posting-time and hashtag notes": "可选的发布时间与 hashtag 备注",
    "Posting-time or hashtag notes": "发布时间或 hashtag 备注",
    "Creator Formula Distillation": "创作者公式蒸馏报告",
    "repeatable creator formula and adaptation rules": "可重复创作者公式与改编规则",
    "Several top creator videos with transcript and performance notes": "多条高表现创作者视频、转写稿与表现备注",
    "US": "美国",
    "Account overview": "账号概览",
    "High vs low interaction comparison": "高低互动对比",
    "Repeatable formulas": "可重复公式",
    "Non-transferable advantages": "不可迁移优势",
    "New-script bridge": "新脚本桥接",
    "Comments from at least 2 products": "至少 2 个商品的评论",
    "Comments from 2+ products": "来自 2 个以上商品的评论",
    "Market": "市场",
    "Product positioning goal": "产品定位目标",
    "Optional price-band notes": "可选的价格带备注",
    "Purchase factor synthesis": "购买因素提炼",
    "Praise keyword synthesis": "好评关键词提炼",
    "Complaint pain-point synthesis": "差评痛点提炼",
    "Price-band difference view": "价格带差异视图",
    "Persona summary": "人群画像总结",
    "Selection and content implications": "选品与内容启发",
    "Real TikTok capture-pack import from ": "真实 TikTok capture-pack 导入自 ",
    " for uncategorized lane. Current board size: ": "，当前用于未分类赛道。当前看板规模：",
    " ranked / ": " 条已排序 / ",
    " qualified.": " 条达标。",
    "qualified with min-like threshold ": "条达标，最低点赞阈值 ",
    "Source profile: ": "来源账号：",
    "session quality: ": "会话质量：",
    "queries: ": "查询词：",
    "topics: ": "主题：",
    "Normalize all candidates into one market before ranking.": "先把所有候选样本统一到同一市场，再进行排序。",
    "Tag each selected video by best reuse purpose: hook, proof, structure, or style.": "给每条入选视频标注最佳复用用途：hook、证明、结构或风格。",
    "Keep the rejected pool so later ranking logic can be improved.": "保留未入选池，便于后续迭代排序逻辑。",
    "Collecting links without enough hook or proof notes for later teardown.": "只收集链接，却没有为后续拆解保留足够的 hook 或证明笔记。",
    "Keep comments grouped by product before merging category-level signals.": "合并品类信号前，先按商品维度保留评论分组。",
    "Prefer repeated user phrases over abstract sentiment summaries.": "优先保留重复出现的用户原话，而不是抽象情绪总结。",
    "Source Product Summary": "来源商品概览",
    "Comment Signal Clusters": "评论信号聚类",
    "Style": "风格",
    "Environment": "环境",
    "Camera": "镜头",
    "Lighting": "灯光",
    "Character": "角色",
    "Product direction": "产品方向",
    "Offer / positioning": "报价 / 定位",
    "Script language": "脚本语言",
    "Proof content": "证明内容",
    "Hook": "钩子",
    "hook": "钩子",
    "Proof": "证明",
    "Close": "收口",
    "CTA": "CTA",
    "Brief": "制作简报",
    "Handoff": "交接",
    "Tone & Pacing": "语气与节奏",
    "Background Sound": "背景声音",
    "Transition / Editing": "转场 / 剪辑",
    "Visual style": "视觉风格",
    "Video type": "视频类型",
    "Posting time": "发布时间",
    "Hashtag habit": "Hashtag 习惯",
    "BGM / audio": "BGM / 音频",
    "Trust-building": "信任建立",
    "Conversion move": "转化动作",
    "Formula Library": "公式库",
    "New-Script Bridge": "新脚本桥接",
    "Account Overview": "账号概览",
    "High Vs Low Interaction Comparison": "高低互动对比",
    "Opening hook": "开头钩子",
    "Conversion rhythm": "转化节奏",
    "BGM / audio mood": "BGM / 音频氛围",
    "Subtitle style": "字幕样式",
    "Transition rhythm": "转场节奏",
    "Silence / pause usage": "留白 / 停顿使用",
    "Generator / Editor Handoff": "生成器 / 剪辑交接",
    "Shot-Level Breakdown": "镜头级拆解",
    "Inferred Original Brief Schema": "反推原始制作简报结构",
    "Hero cue / hook frame": "首屏关键线索 / 钩子画面",
    "Support frame or subtitle": "辅助画面或字幕",
    "Continuation CTA": "延续式 CTA",
    "continuation close": "延续式收口",
    "editorial / social-native": "编辑感 / 原生社交风格",
    "packaging study": "包装研究",
    "download enrichment": "下载补全",
    "portable format": "可迁移格式",
    "General TikTok packaging study": "通用 TikTok 包装研究",
    "Preserve the recognition-first packaging and rewrite the proof layer with owned assets.": "保留先识别后证明的包装顺序，并用自有资产重写证明层。",
    "Recovered cover image URL that can support frame review or first-frame comparison.": "补全出的封面图 URL，可用于逐帧复核或首帧对比。",
    "Cover / key frame source": "封面 / 关键帧来源",
    "creator-top-1": "创作者高互动样本 1",
    "creator-low-1": "创作者低互动样本 1",
    "paste-video-link": "待补视频链接",
    "paste-screenshot-path-or-link": "待补截图路径或链接",
    "High-interaction sample used as the positive contrast case.": "高互动样本，作为正向对照案例。",
    "Lower-interaction sample used to isolate what changed.": "低互动样本，用于隔离变化因素。",
    "High-interaction comparison": "高互动对比",
    "Low-interaction comparison": "低互动对比",
    "Opening recognition cue": "开场识别线索",
    "Setup / context beat": "铺垫 / 场景承接节拍",
    "Compressed setup, early proof, soft close": "压缩铺垫，尽早给证明，轻量收口",
    "Social-native editorial context with one dominant recognition cue": "带一个主识别线索的原生社交编辑环境",
    "Short premise-led beats with quick recognition framing": "以核心前提驱动的短节拍，并快速给出识别画面",
    "Visible and believable over polished and cinematic": "优先可见、可信，而不是过度精致或电影化",
    "Original sound or recovered social-native audio": "原始音频或补全后的原生社交音频",
    "Fast cuts or compact beat changes": "快速切换或紧凑节拍变化",
    "Return to the main cue and guide the next click or watch": "回到主线索，并引导下一次点击或继续观看",
    "Recognition-first opening and compressed premise": "先识别、后展开的开场与压缩前提",
    "Soft continuation close": "轻量延续式收口",
    "Use available talent, object, or demo footage only": "只使用现成的人物、物件或演示素材",
    "Talent / scene availability": "人物 / 场景可用性",
    "Hard-sell CTA may break the native pacing": "硬卖式 CTA 可能破坏原生节奏",
    "profile=": "账号=",
    "session=": "会话=",
    "ranked=": "已排序=",
    "qualified=": "达标=",
    "min_likes=": "最低点赞阈值=",
    "likes=": "点赞=",
    "comments=": "评论=",
    "shares=": "分享=",
    "plays=": "播放=",
        "One keyword": "1 个关键词",
        "At least 5 candidate videos, links, or screenshots": "至少 5 条候选视频、链接或截图",
        "One video link or storyboard summary": "1 条视频链接或一份分镜摘要",
        "One video or visual summary.": "1 条视频或一份视觉摘要。",
        "One video or visual summary": "1 条视频或一份视觉摘要",
        "15-30 candidates with basic metrics": "15-30 条带基础指标的候选样本",
    "Search-result screenshots": "搜索结果截图",
    "Market and audience note": "市场与受众备注",
    "Market context": "市场上下文",
    "and audience note": "与受众备注",
    "yes": "是",
    "recommended": "建议",
    "15 TikTok search results with links, metrics, and first-hook notes": "15 条 TikTok 搜索结果，含链接、基础指标与首拍钩子笔记",
    "Comments from 3 products with repeated phrases highlighted": "来自 3 个商品的评论样本，已标出重复短语与购买型语言",
    "2+ competitor accounts with one week's posts and performance notes": "2 个以上竞品账号的一周帖子与表现备注",
    "Recent post table with metrics, hooks, and content-type labels": "最近帖子表，包含指标、钩子和内容类型标签",
    "Recent post table with metrics, 钩子s, and content-type labels": "最近帖子表，包含指标、钩子和内容类型标签",
    "Primary keyword or product phrase": "核心关键词或产品短语",
    "Target market": "目标市场",
    "Target audience": "目标人群",
    "Date window or freshness requirement": "发布时间窗口或时效要求",
    "Sort order": "排序方式",
    "Whether to keep only TikTok Shop cart videos": "是否仅保留 TikTok Shop 购物车视频",
    "Do not rank on views alone. Keep reuse value in the scoring logic.": "不要只按播放量排序，排序逻辑里必须保留复用价值判断。",
    "If live browsing is unavailable, rely on user-provided screenshots, exports, or copied links.": "如果当前无法实时浏览，就依赖用户提供的截图、导出文件或复制链接。",
    "Ranked shortlist": "排序短名单",
    "Structured collection board": "结构化采集主表",
    "Reason each selected video matters": "每条入选视频为什么值得研究",
    "Scene-03 shortlist handoff": "Scene 03 交接短名单",
    "Scene 01 Spot Check": "场景 01 成品质检",
    "Scene 17 Spot Check": "场景 17 成品质检",
    "Study-next recommendation": "下一步优先深拆建议",
    "TikTok-native ranked-pattern conclusions": "TikTok 原生排序模式结论",
    "Reusable adaptation rules grounded in the capture pack": "基于 capture pack 的可复用改编规则",
    "Candidate set is from one market": "候选样本来自同一市场",
    "At least basic performance signals exist": "至少具备基础表现信号",
    "Useful-for tags can be assigned": "可以给样本打上用途标签",
    "Publish window and sort rule are explicit": "发布时间窗口与排序规则已明确",
    "Real TikTok anonymous-session capture. Comment sampling is missing in this pack.": "当前为真实 TikTok 匿名会话采集包；本包暂不含评论采样。",
    "This capture pack already contains a usable shortlist of TikTok posts worth deeper study because the strongest rows now carry recoverable caption, hook, and topic signals in addition to ranking metrics.": "这份 capture pack 已经具备可直接深拆的 TikTok 短名单，因为最强样本除了排序指标外，还带有可恢复的 caption、hook 和主题信号。",
    "The board can now function as a real reusable intake layer because the shortlist is no longer only score-based; it preserves more of the source caption and packaging logic for teardown.": "这份看板已经能承担真正可复用的 intake 层，因为 shortlist 不再只看分数，而是保留了更多源 caption 和包装逻辑，便于继续拆解。",
    "Take the top three into deeper teardown immediately and assign each one a clear study lane before anyone starts analyzing ad hoc.": "先把 Top 3 立即送去深拆，并在零散分析前先给每条样本明确研究方向。",
    "Lock publish-time window, region, sort rule, and cart-video scope before collection.": "采集前先锁定发布时间窗口、地区、排序规则与购物车视频范围。",
    "Ranking on views only and ignoring reuse value.": "只按播放量排序，忽略复用价值。",
    "Mixing organic-looking viral hits with commerce-heavy posts without labeling the difference.": "把偏自然流量的爆款和强带货内容混在一起，却没标清差异。",
    "Mixing multiple markets or product intents in one shortlist.": "在同一 shortlist 里混入多个市场或多个产品意图。",
    "Run scene 01 to collect and rank the best viral videos for one keyword and market. Treat publish-time window, region, sort order, and whether to include only TikTok Shop cart videos as explicit inputs, score for reuse value instead of views alone, and finish with a shortlist that can move directly into scene 03.": "按场景 01 执行：围绕一个关键词或品类和单一市场，采集并排序最值得研究的爆款视频。把发布时间窗口、地区、排序方式和是否只看带购物车视频作为明确输入，不要只按播放量排序，而是按复用价值排序，最后给出可直接进入场景 03 的 shortlist。",
    "The final deliverable must include: ranked shortlist, Structured collection board, reason each selected video matters, Scene-03 shortlist handoff.": "最终必须产出：排序短名单、结构化采集主表、每条入选视频为什么值得研究、Scene 03 交接短名单。",
    "The final deliverable must include the following operator-ready outputs: ranked shortlist, Structured collection board, reason each selected video matters, Scene-03 shortlist handoff.": "最终必须产出以下可直接给运营使用的结果：排序短名单、结构化采集主表、每条入选视频为什么值得研究、Scene 03 交接短名单。",
    "Rank the candidate pool before doing any deeper analysis.": "先完成候选池排序，再进入任何深层分析。",
    "Tag each shortlisted video by best reuse purpose: hook, proof, structure, or style.": "给每条 shortlist 视频标注最适合复用的用途：hook、证明、结构或风格。",
    "Preserve shopping intent and TikTok Shop signals as separate fields, not buried inside notes.": "把购买意图与 TikTok Shop 信号单独列出，不要埋在备注里。",
    "ranked shortlist plus teardown priority": "排序短名单与深拆优先级",
    "pain-language synthesis and persona guidance": "买家语言提炼与人群画像启发",
    "weekly competitor report and action board": "竞品周报与动作看板",
    "performance retro and next-cycle test plan": "表现复盘与下轮测试计划",
    "The shortlist is ranked and limited to the strongest candidates.": "短名单已经排序，并限制在最强样本内。",
    "Each selected video has a concrete why-selected reason.": "每条入选视频都有明确的入选原因。",
    "Commerce confidence and reuse purpose are explicit per shortlisted video.": "每条 shortlisted 视频都明确写出商业置信度与复用用途。",
    "The operator knows which videos should move into the next teardown workflow.": "操作者知道哪些视频应该立即进入下一步深拆流程。",
    "Selected for market traction": "因市场牵引力入选",
    "Selected for market traction, download enrichment, portable format": "因市场牵引力、下载补全和可迁移格式入选",
    "Selected for comment density": "因评论密度入选",
    "supporting traction:": "支撑信号：",
    "next move:": "下一步动作：",
    "Route to": "优先进入",
    "Top candidate topic:": "头部候选主题：",
    "Top candidate hook:": "头部候选钩子：",
    "These are not just top-view posts; they are ranked candidates with reusable packaging traits and stronger recovered text evidence.": "这些不只是高播放帖子，而是带有可复用包装特征和更强文本恢复证据的排序候选样本。",
    "The best shortlist items should be routed by reuse value, not only raw numbers.": "最值得进入下一步的样本应按复用价值分流，而不只是看原始数据大小。",
    "Build the shortlist table first.": "先搭好短名单主表。",
    "Explain why each selected item deserves operator attention.": "解释为什么每条入选视频值得被优先关注。",
    "Define the minimum schema for future collection rounds.": "定义下一轮采集至少要补齐哪些字段。",
    "Recommend what to do immediately after collection.": "明确采集完成后立刻要做什么。",
    "Video link": "视频链接",
    "Traceability into later teardown": "便于后续深拆回溯",
    "Traceability": "便于追溯到后续深拆",
    "Region": "地区",
    "Keep one market scope": "确保只看单一市场",
    "Sort by": "排序方式",
    "Reproducible collection pass": "保证采集轮次可复现",
    "Cart / shop presence": "购物车 / 店铺信号",
    "Commerce intent": "判断商业意图",
    "Post date": "发布时间",
    "Freshness": "判断新鲜度",
    "Views / likes / comments": "播放 / 点赞 / 评论",
    "Basic performance": "保留基础表现形态",
    "Hook summary": "钩子摘要",
    "Later breakdown": "供后续拆解使用",
    "Useful-for tag": "复用用途标签",
    "Routing to next workflow": "决定进入哪条下游流程",
    "Summary": "汇总",
    "Profile summary": "账号汇总",
    "Ranked video": "排序视频",
    "Aggregate markdown report from the real TikTok capture pack.": "真实 TikTok capture pack 的聚合 Markdown 报告。",
    "If multiple markets are mixed together, split the board before drawing conclusions.": "如果多个市场混在一起，先拆分看板，再下结论。",
    "20-40 comments per product": "每个商品 20-40 条评论",
    "Positioning goal": "定位目标",
    "Price-band context": "价格带上下文",
    "The strongest repeated user language in this TikTok comment pack now resolves into clearer purchase, trust, and complaint clusters, with reply-chain pressure separating shallow reaction from real objection handling.": "这份 TikTok 评论包里最强的重复用户语言，已经能沉淀成更清晰的购买、信任与抱怨聚类，且回复链压力可以把浅层反应和真实异议处理分开。",
    "This matters because the operator now has cleaner buyer-language clusters, source-product labels, duplicate-collapsed quotes, and reply-chain cues instead of one flat pile of comments.": "这很重要，因为操作者现在拿到的不再是一堆平铺评论，而是更干净的买家语言聚类、来源商品标签、去重后的高频原话和回复链线索。",
    "Use the cleaned complaint and trust clusters to write moderator replies, content FAQ lines, and positioning copy before the next test cycle.": "下一轮测试前，直接用清洗后的抱怨与信任聚类去写评论回复、内容 FAQ 话术和定位文案。",
    "High-performing videos": "高表现视频",
    "Transcripts": "转写稿",
    "transcripts": "转写稿",
    "Creator playbook": "创作者打法手册",
    "Adaptation path": "改编路径",
    "Creator Brief Reconstruction": "创作者制作简报重建",
    "Creator Brief": "创作者制作简报",
    "制作 brief": "制作简报",
    "制作 Brief": "制作简报",
    "反向推断这条视频背后的提示词或制作 brief，把创作意图拆成视觉、镜头、旁白、节奏模块，并把低置信度猜测标出来。": "反向推断这条视频背后的提示词或制作简报，把创作意图拆成视觉、镜头、旁白、节奏模块，并把低置信度猜测标出来。",
    "反推原始 Brief / 产品适配 Brief": "反推原始制作简报 / 产品适配制作简报",
    "Pattern Area": "模式维度",
    "What Repeats": "重复点",
    "Example Evidence": "示例证据",
    "Transferable Pattern": "可迁移模式",
    "Why It Transfers": "为什么可迁移",
    "How To Adapt": "如何改写",
    "Visual rhythm": "视觉节奏",
    "Proof style": "证明方式",
    "Missing comments": "缺少评论样本",
    "Lane": "处理通道",
    "What To Do Now": "当前动作",
    "Why This Lane Exists": "设这个通道的原因",
    "Profile session quality:": "账号会话质量：",
    "Layer": "层级",
    "Pacing": "节奏",
    "Test": "测试",
    "Editing / pacing": "剪辑 / 节奏",
    "Owned-product rewrite": "自有产品改写",
    "Prompt hardening": "提示词加固",
    "Need user's product, proof object, or collaborator context": "需要用户产品、证明物或合作方语境",
    "Operator / strategist": "运营 / 策略",
    "Need final platform, asset, and voice constraints": "需要最终投放平台、素材条件与口播约束",
    "Creative lead": "创意负责人",
    "Human-, context-, or outcome-led first frame with compressed topic cue": "以人物、语境或结果为首屏主体，并压缩主题线索。",
    "Open on the strongest recognition cue, add one proof beat, then close with a continuation move": "先给最强的识别线索，再补一个证明节拍，最后用延续式动作收口。",
    "Minimal explanation, premise support, soft continuation CTA": "少解释、支撑前提、轻量延续式 CTA。",
    "Short beats, little dead air, no long setup before the reason to watch is clear": "短节拍、少空镜、在看点明确前不做长铺垫。",
    "Source-account authority may be doing part of the work, so the rewritten brief needs a stronger owned proof object.": "源账号权威可能承担了部分效果，所以重写后的制作简报需要更强的自有证明物。",
    "Low-confidence inference: the exact camera prompt or edit stack is not recoverable from ranked post metadata alone.": "低置信度推断：仅靠排序帖元数据，无法恢复精确的镜头提示词或剪辑栈。",
    "If the user product lacks a fast visual cue, the inferred pacing may still fail after adaptation.": "如果用户产品缺少快速可识别的视觉线索，即使完成改写，原节奏也可能失效。",
    "The inferred brief works because it does not ask the viewer to decode too much before the premise is obvious.": "这份反推制作简报之所以有效，是因为它不会在前提还没清楚前，就让观众承担过多理解成本。",
    "The likely hidden system prompt is not 'make it viral'. It is 'make the first frame legible, then support it with one proof device and one continuation move.'": "更像的隐藏系统提示不是“做成爆款”，而是“先让首屏可读，再用一个证明装置和一个延续动作把它托起来”。",
    "Keep the same recognition-first structure but replace the source proof cue": "保留同样的识别优先结构，但把源视频里的证明线索换掉",
    "Turn this inferred brief into a render- or shoot-ready block format": "把这份反推制作简报加固成可直接渲染或拍摄的模块格式",
    "A large platform account has built-in distribution and credibility most projects do not have.": "大平台账号自带分发和公信力，这不是大多数项目天然具备的条件。",
    "Recognition from known creators or cultural moments may be doing part of the ranking work.": "来自知名创作者或文化瞬间的识别度，可能承担了部分排序优势。",
    "Without sampled comments, true audience-language resonance is under-evidenced in this import.": "由于缺少评论采样，这次导入对真实受众语言共振的证据仍然偏弱。",
    "Keep tracking which winners depend on featured-person recognition or official-account lift.": "继续追踪哪些赢家样本依赖人物识别度或官方账号 lift。",
    "This stops the team from mistaking distribution advantage for portable format logic.": "这样可以避免团队把分发优势误认成可迁移的格式逻辑。",
    "Run one creator-led rewrite and one proof-object-led rewrite from the top account references.": "基于头部账号参考，分别跑一版创作者驱动改写和一版证明物驱动改写。",
    "This is the fastest way to learn whether the pattern survives after authority replacement.": "这是验证“替换掉权威之后，这套模式是否仍成立”的最快方法。",
    "Do not escalate rows whose only real edge is platform-scale authority.": "不要升级那些唯一优势只是平台级权威的样本行。",
    "Those rows distort smaller-account replication planning.": "这类样本会扭曲小账号复刻规划。",
    "Open with immediate recognition or emotional clarity": "以即时识别或情绪清晰度开头",
    "Recognition compresses decision time on TikTok": "识别感能压缩用户在 TikTok 上的决策时间",
    "Swap in a figure, object, or cue your audience already cares about": "替换成受众本来就关心的人物、物件或首屏线索",
    "Stay short and premise-led": "保持短、前提先行",
    "The format works because it does not over-teach": "这种格式有效，是因为它不会过度教学",
    "Strip excess setup before the main cue lands": "在主线索落地前，删掉多余铺垫",
    "Borrow trust from the account, featured talent, or event context": "从账号、人物或事件语境借来信任",
    "Trust can be transferred via stronger proof objects": "信任可以通过更强的证明物转移",
    "Use receipts, social proof, or known collaborators if account authority is weaker": "如果账号权威更弱，就用凭证、社交证明或已知合作方补足",
    "Use continuation energy instead of hard closing": "用延续感代替硬收口",
    "Soft progression fits social-native viewing better": "轻推进更符合社交原生观看",
    "Route toward next watch, next profile action, or soft save/share": "把动作导向下一次观看、下一步主页动作或轻量收藏 / 分享",
    "This TikTok account sample suggests a repeatable editorial formula: attach the post to a recognizable creator, story, or cultural moment, then use minimal copy to let affinity do the work.": "这组 TikTok 账号样本显示出一套可重复的编辑型公式：先把内容挂到可识别的创作者、故事或文化瞬间上，再用极少文案让亲和感自行发挥作用。",
    "The pattern is useful for TikTok projects that need stronger social-native packaging without long explanation-heavy intros.": "这种模式适合需要更强社交原生包装、但又不想靠长解释开头的 TikTok 项目。",
    "Translate the account's strongest editorial packaging moves into a reusable creator- or community-led content brief.": "把这个账号最强的编辑型包装动作，翻译成可复用的创作者型或社群型内容制作简报。",
    "This real TikTok reference suggests a prompt or production brief built around fast recognition, minimal explanation, and one trust-bearing social cue rather than heavy narrative complexity.": "这条真实 TikTok 参考视频反推出的是一份以快速识别、少解释、单个信任支点为核心的提示词或制作简报，而不是重叙事、重世界观的复杂脚本。",
    "That makes the inferred brief reusable: the operator can preserve visual pacing and premise order while swapping the proof object or featured cue onto a different product or account.": "这让反推出来的制作简报具备复用价值：保留视觉节奏和前提顺序后，只要替换证明物或主线线索，就能迁移到别的产品或账号。",
    "Treat the inferred brief as a structured creation blueprint, then mark which parts depend on source-account authority versus portable shot and copy logic.": "把这份反推制作简报当成结构化创作蓝图使用，并明确区分哪些部分依赖源账号权威，哪些部分属于可迁移的镜头与文案逻辑。",
    "The strongest single-video breakdown target in this real TikTok pack wins by making the first frame instantly legible, then using authority, context, or a featured-person cue as compressed proof.": "这条真实 TikTok 单视频样本之所以能跑出来，核心在于首屏先让人秒懂，再用权威、语境或人物线索做压缩证明。",
    "This matters because the reusable asset is not surface polish. It is the sequence that moves from recognition to proof to a soft continuation close without over-explaining.": "真正可复用的资产不是表层精致感，而是从识别、到证明、再到轻量延续收口的顺序，并且全程不过度解释。",
    "Rebuild the reference in order, then rewrite the proof layer so an owned product, creator, or evidence object can carry the same decision logic.": "先按顺序重建参考视频，再把证明层改写成自有产品、创作者或证据物也能承接的同一套决策逻辑。",
    "Source account baseline:": "来源账号基线：",
    "Recovered hook:": "已恢复钩子：",
    "Recovered 钩子:": "已恢复钩子：",
    "Recovered topic cue:": "已恢复主题线索：",
    "Authority signal:": "权威信号：",
    "The real mechanism is recognition-first compression: the viewer understands who or what matters before the video spends attention on explanation.": "真正起作用的机制，是先做识别压缩：观众会在视频花时间解释前，先明白谁或什么才是重点。",
    "The proof layer works because the source account can borrow trust from authority, a familiar creator, or a culturally legible moment instead of spelling out every claim.": "这条视频的证明层之所以成立，是因为源账号可以借用权威、熟悉创作者或文化线索来转移信任，而不是把每个论点都解释一遍。",
    "Portable logic: first-frame clarity plus compressed proof.": "可迁移逻辑：首屏清晰 + 压缩证明。",
    "Non-portable lift: official-account authority, featured-talent recognition, or distribution advantage.": "不可直接迁移的加成：官方账号权威、人物识别度，或平台分发优势。",
    "Keep the hook structure and proof order": "保留钩子结构与证明顺序",
    "Keep only the recognition-first shell": "只保留识别优先的外壳",
    "Reframe the topic, talent, and close for a new product or creator": "围绕新的产品或创作者重写主题、人物线索与收口方式",
    "May lose the original trust transfer if the new cue is not instantly legible": "如果新的线索不够一眼可懂，原本的信任转移效果可能会丢失",
    "Lead candidate:": "头部候选：",
    "Qualified control:": "达标对照：",
    "The account does not need to over-explain. It packages a familiar person or cultural cue and relies on fast recognition plus account trust.": "这个账号不需要过度解释。它更像是在包装一个熟悉的人物或文化线索，再用快速识别和账号信任把内容推过去。",
    "The transferable lesson is not 'be TikTok'. It is to reduce friction between first-frame recognition and the emotional reason to keep watching.": "真正可迁移的经验不是“变得像 TikTok”，而是降低首屏识别和继续观看动机之间的摩擦。",
    "Map repeated 钩子, pacing, proof, and CTA patterns separately.": "把重复出现的钩子、节奏、证明和 CTA 模式分开拆。",
    "Run scene 17 to distill one creator's repeatable content formula across multiple videos. Map repeated 钩子, pacing, proof, and CTA patterns, then separate transferable logic from creator-specific advantage.": "按场景 17 执行：从同一创作者的多条视频里蒸馏可重复的内容公式。把重复出现的钩子、节奏、证明和 CTA 模式分开拆，再区分可迁移逻辑和创作者专属优势。",
    "Run scene 05 to reverse-engineer the likely prompt or production brief behind one video. Infer the original creative brief, map it into generator-ready blocks such as Style, Environment, Camera, Lighting, Character, and shot-level structure, then produce a product-adapted version with field-level low-confidence labels.": "按场景 05 执行：反推 1 条视频背后的提示词或制作简报。先推断原始创意意图，再拆成可直接交给生成器的模块，如风格、环境、镜头、灯光、角色与分镜结构，最后产出带字段级低置信度标记的产品适配版。",
    "Run scene 04 to fully break down one TikTok or Douyin video. Rebuild it beat by beat with a timeline table, BGM read, 钩子-to-conversion rhythm, and video-type classification, then separate the real mechanism from surface style and recommend one concrete adaptation path.": "按场景 04 执行：完整拆解 1 条 TikTok 或抖音视频。用时间轴逐拍重建画面、BGM、钩子到转化节奏与视频类型，再把真正有效的机制与表层风格分开，最后给出 1 条具体可执行的改编路径。",
    "Translate observed output into generator-ready prompt blocks such as Style, Environment, Tone & Pacing, Camera, Lighting, Character, Shots, Background Sound, and Transition / Editing.": "把观察到的成品翻译成可直接喂给生成器的提示块，例如风格、环境、语气与节奏、镜头、灯光、角色、分镜、背景声音和转场 / 剪辑。",
    "Translate observed output into generator-ready prompt blocks such as Style, Environment, Camera, Lighting, Character, and shot-level structure.": "把观察到的成品翻译成可直接喂给生成器的提示块，例如风格、环境、镜头、灯光、角色与分镜结构。",
    "Split the deliverable into inferred-original and product-adapted variants.": "交付结果必须拆成“反推原版”和“产品适配版”两层。",
    "If evidence is thin, keep low-confidence labels visible in the final brief.": "如果证据偏薄，最终制作简报里要保留低置信度标记。",
    "Use several creator samples before declaring a formula.": "在宣称形成公式前，先使用多个创作者样本。",
    "Do not confuse admiration for the creator with reusable production rules.": "不要把对创作者的赞赏误当成可复用的制作规则。",
    "Rebuild the video in sequence before drawing any conclusions.": "先按顺序重建视频，再下结论。",
    "Explicitly separate creator-specific polish from transferable conversion logic.": "明确区分创作者专属包装与可迁移的转化逻辑。",
    "Support both spoken-script and no-voiceover videos by capturing subtitles, gestures, and motion-based proof when needed.": "同时支持有口播与无口播视频，必要时用字幕、动作和运动证明链补齐分析。",
    "The hook, proof, and close are reconstructed in order.": "钩子、证明与收口已按顺序复原。",
    "Role": "作用",
    "Adaptation Guardrail": "改编护栏",
    "Safer": "更稳妥",
    "More aggressive": "更激进",
    "Hook logic": "钩子逻辑",
    "Proof logic": "证明逻辑",
    "CTA style": "CTA 风格",
    "Layer": "层级",
    "Visual direction": "视觉方向",
    "Shot plan": "分镜计划",
    "Voiceover logic": "口播逻辑",
    "Shot language": "镜头语言",
    "Narrative pacing": "叙事节奏",
    "Pacing design": "节奏设计",
    "Hook formula": "钩子公式",
    "Official-account authority": "官方账号权威",
    "Featured-talent lift": "人物识别度加成",
    "Creator-Specific Advantage": "创作者专属优势",
    "Operator Dispatch": "操作分发",
    "Watch": "持续观察",
    "Suppress": "暂不升级",
    "Build one smaller-account version with a stronger owned proof object before filming anything else.": "先做一版更适合小账号的版本，并在开拍前补上更强的自有证明物。",
    "Measure whether lighter copy plus faster recognition improves early engagement enough to justify reusing the formula.": "验证更轻的文案配合更快的识别，是否足以提升早期互动，并支撑这套公式继续复用。",
    "Translate observed output into generator-ready prompt blocks such as Style, Environment, 语气与节奏, Camera, Lighting, Character, Shots, 背景声音, and 转场 / 剪辑.": "把观察到的成品翻译成可直接喂给生成器的提示块，例如风格、环境、语气与节奏、镜头、灯光、角色、分镜、背景声音和转场 / 剪辑。",
    "反推原始 Brief": "反推原始制作简报",
    "产品适配 Brief": "产品适配制作简报",
    "反推提示词与 Brief": "反推提示词与制作简报",
    "for user product is possible": "已具备面向用户产品的改写路径",
    "转写稿s": "转写稿",
    "Profile 会话质量": "账号会话质量",
    "Keep the output socially legible before detail appears": "在细节出现前先保证成品具备社交平台可读性",
    "Lead with the cue that makes the viewer instantly care": "先给出让用户立刻在意的线索",
    "Avoid explainers that delay the core promise": "避免长解释拖慢核心承诺露出",
    "Support the visible cue instead of overpowering it with exposition": "用口播去支撑可见线索，而不是用解释把它盖过去",
    "Human-first, emotion-first, or culture-first caption packaging": "以人物优先、情绪优先或文化语境优先的 caption 包装开场",
    "Likely dependent on featured creator/performance clip rather than explanation-led structure": "更依赖出镜创作者或表演片段，而不是解释驱动结构",
    "Trust rides on official account authority, featured people, and recognizable context": "信任主要建立在官方账号权威、出镜人物和可识别语境上",
    "Soft teaser or continuation toward more content": "用柔性预告或续看方式引向更多内容",
    "Assign one top-ranked post to a creator-led rewrite lane and one to a moment/outcome rewrite lane.": "把一条高排名帖子分到创作者驱动改写线，另一条分到时刻 / 结果驱动改写线。",
    "Describe why the creator's pattern works.": "说明这位创作者的模式为什么有效。",
    "Describe how to migrate the pattern to the user's product.": "说明如何把这套模式迁移到用户自己的产品。",
    "Describe how to migrate the pattern to the user's product or account.": "说明如何把这套模式迁移到用户自己的产品或账号。",
    "2+ competitor accounts": "2 个以上竞品账号",
    "One weekly batch of posts": "1 周帖子批次",
    "Prior week notes": "上周备注",
    "Per-post performance context": "逐帖表现上下文",
    "Target market": "目标市场",
    "3-5 accounts in one matrix": "1 个矩阵里包含 3-5 个账号",
    "Per-account weekly summary": "分账号周度总结",
    "Cross-account comparison": "跨账号横向对比",
    "Notable shifts": "关键变化",
    "Strategy-shift view": "策略变化视角",
    "Implications for the user": "对用户的影响",
    "Posts are grouped by account and week": "帖子已按账号与周维度分组",
    "Shift vs prior week can be stated": "可以明确写出相对上周的变化",
    "Weekly response actions can be prioritized": "可以给每周响应动作排优先级",
    "Cross-account comparison is possible": "可以进行跨账号横向对比",
    "If only one week exists, mark it as baseline rather than trend.": "如果目前只有 1 周数据，应标记为基线，而不是趋势判断。",
    "Group posts by account and week before comparing anything.": "开始比较前，先按账号和周维度整理帖子。",
    "Highlight weekly shifts, not just weekly totals.": "重点标出周度变化，而不是只列周度总量。",
    "Compare accounts horizontally, not as isolated mini-reports.": "横向比较账号，而不是拆成互不相干的小报告。",
    "Finish with actions the user should take this week.": "最后必须落到本周用户该采取的动作。",
    "This TikTok capture pack establishes a usable weekly competitor-account baseline: the account is winning with a small number of editorially packaged, emotion-first or culture-first posts.": "这份 TikTok capture pack 已经建立了可用的竞品账号周基线：该账号当前主要靠少量、编辑包装感强、情绪优先或文化优先的帖子取胜。",
    "Even one weekly baseline is enough to decide what kind of post packaging deserves continued tracking versus what is just account noise.": "即使只有 1 周基线，也足够判断哪些帖子包装方式值得继续追踪，哪些只是账号噪声。",
    "Use this pack as the baseline week, then compare the next capture against the same account fields to spot packaging or performance shifts.": "先把这份包当成基线周，下一次采集再按相同账号字段对比，就能识别包装或表现的变化。",
    "Recent post list": "最近帖子列表",
    "Some performance signal per post": "每条帖子至少有部分表现信号",
    "Views, likes, comments, saves, shares": "播放、点赞、评论、收藏、分享",
    "Hook / title notes": "hook / 标题备注",
    "Content-type labels": "内容类型标签",
    "Optional conversion or ROI context": "可选的转化或 ROI 上下文",
    "High vs low performance grouping": "高低表现分组",
    "Performance pattern summary": "表现模式总结",
    "Winning traits": "高表现特征",
    "Losing traits": "低表现特征",
    "Next-cycle plan": "下一轮计划",
    "Cluster posts by pattern, not just by publish date.": "帖子要按内容模式聚类，而不是只按发布时间排列。",
    "Compare high-performing and low-performing groups explicitly.": "明确比较高表现组与低表现组。",
    "Write explicit do-more, do-less, and stop rules.": "明确写出多做、少做和停止的规则。",
    "Turn the retro into one next-cycle testing plan.": "把复盘结果落成 1 份下一轮测试计划。",
    "Within this TikTok account sample, the likely winning pattern is short, editorially framed posts that attach to a recognizable person, story, or moment instead of leading with heavy explanation.": "在这组 TikTok 账号样本里，更可能取胜的是短而有编辑框架的帖子，它们会挂靠到可识别的人、故事或瞬间，而不是一开始就长篇解释。",
    "This is useful as a retro template because it converts raw ranked-post data into do-more, do-less, and next-test rules for the next cycle.": "它适合作为复盘模板，因为能把原始排序帖数据转成下一轮的多做、少做和测试规则。",
    "Cluster the next account batch around people-led, moment-led, and explanation-led posts to confirm which packaging family deserves more volume.": "下一批账号内容应围绕人物驱动、瞬间驱动和解释驱动三类帖子聚类，以确认哪一类包装方式值得加量。",
    "If metrics are incomplete, keep weak conclusions explicitly labeled.": "如果指标不完整，必须把弱结论显式标记出来。",
    "Adaptation path for user product is possible": "已经具备面向用户产品的改编路径",
    "Posts can be clustered by pattern": "帖子可以按模式聚类",
    "Winners and losers are distinguishable": "高低表现内容可以明确区分",
}

SHARED_TEMPLATE_TEXT_PATTERN = re.compile(
    "|".join(re.escape(key) for key in sorted(SHARED_TEMPLATE_TEXT_REPLACEMENTS, key=len, reverse=True))
)

EXACT_DISPLAY_VALUE_REPLACEMENTS = {
    "medium": "中",
    "low": "低",
    "high": "高",
    "low-to-medium": "低到中",
    "high-to-medium": "高到中",
}


def localize_template_text(value: str) -> str:
    text = display_clean_text(value)
    if not text:
        return ""
    text = SHARED_TEMPLATE_TEXT_PATTERN.sub(lambda match: SHARED_TEMPLATE_TEXT_REPLACEMENTS[match.group(0)], text)
    text = re.sub(r"(?<=\d)\s+likes(?=,|\b)", " 点赞", text)
    text = re.sub(r"(?<=\d)\s+plays(?=,|\b)", " 播放", text)
    text = re.sub(r"(?<=\d)\s+shares(?=,|\b)", " 分享", text)
    text = re.sub(r"(?<=\d)\s+comments(?=,|\b)", " 评论", text)
    text = text.replace("session quality: browser_same_origin_api_ok", "会话质量：浏览器同源接口正常")
    text = text.replace("session quality: tikmatrix_profile_posts_export", "会话质量：TikMatrix 主页帖子导出")
    text = text.replace("session quality: unknown", "会话质量：待补")
    text = text.replace("queries: none", "查询词：未提供")
    text = text.replace("topics: none", "主题：未提供")
    text = text.replace("会话质量：browser_same_origin_api_ok", "会话质量：浏览器同源接口正常")
    text = text.replace("会话质量：tikmatrix_profile_posts_export", "会话质量：TikMatrix 主页帖子导出")
    text = text.replace("会话质量：unknown", "会话质量：待补")
    text = text.replace("查询词：none", "查询词：未提供")
    text = text.replace("主题：none", "主题：未提供")
    return EXACT_DISPLAY_VALUE_REPLACEMENTS.get(text.lower(), text)


def render_localized_markdown(report: dict) -> str:
    markdown_payload = deepcopy(report)
    metadata = markdown_payload.setdefault("metadata", {})
    working_context = markdown_payload.setdefault("working_context", {})
    executive = markdown_payload.setdefault("executive_summary", {})

    metadata["project"] = localize_project_text(metadata.get("project", "")) or metadata.get("project", "")
    metadata["title"] = localized_report_title(metadata)
    metadata["scene_title"] = scene_label_zh(metadata.get("scene", "")) or metadata.get("scene_title", "")
    metadata["deliverable_type"] = normalize_text(metadata.get("deliverable_type", "")) or "TikTok 增长运营交付"
    metadata["status"] = localize_status_text(metadata.get("status", "")) or "草稿"

    executive["conclusion"] = localize_template_text(executive.get("conclusion", ""))
    executive["why_it_matters"] = localize_template_text(executive.get("why_it_matters", ""))
    executive["next_action"] = localize_template_text(executive.get("next_action", ""))
    executive["confidence"] = localize_template_text(executive.get("confidence", ""))

    working_context["summary"] = localize_template_text(working_context.get("summary", ""))
    for key in [
        "inputs",
        "minimum_evidence",
        "ideal_evidence",
        "constraints",
        "requested_outputs",
        "ready_checklist",
    ]:
        working_context[key] = [localize_template_text(item) for item in normalize_string_list(working_context.get(key))]

    operator_guide = markdown_payload.setdefault("operator_guide", {})
    for key in ["operator_checklist", "common_failure_modes"]:
        operator_guide[key] = [localize_template_text(item) for item in normalize_string_list(operator_guide.get(key))]

    execution_template = markdown_payload.setdefault("execution_template", {})
    for key in [
        "recommended_request",
        "recommended_request_zh",
    ]:
        execution_template[key] = localize_template_text(execution_template.get(key, ""))
    for key in [
        "recommended_runner_args",
        "codex_prompt_scaffold",
        "codex_prompt_scaffold_zh",
        "workflow_steps",
        "output_checklist",
    ]:
        execution_template[key] = [localize_template_text(item) for item in normalize_string_list(execution_template.get(key))]
    localized_inputs = []
    for item in execution_template.get("variable_inputs", []) or []:
        if not isinstance(item, dict):
            continue
        localized_inputs.append(
            {
                "name": localize_template_text(item.get("name", "")),
                "meaning": localize_template_text(item.get("meaning", "")),
                "example": localize_template_text(item.get("example", "")),
                "required": localize_template_text(item.get("required", "")),
            }
        )
    execution_template["variable_inputs"] = localized_inputs

    localized_evidence = []
    for item in markdown_payload.get("evidence", []) or []:
        if not isinstance(item, dict):
            continue
        localized_evidence.append(
            {
                "label": localize_template_text(item.get("label", "")),
                "detail": localize_template_text(item.get("detail", "")),
                "source": localize_template_text(item.get("source", "")),
            }
        )
    markdown_payload["evidence"] = localized_evidence

    localized_assets = []
    for item in markdown_payload.get("assets", []) or []:
        if not isinstance(item, dict):
            continue
        localized_assets.append(
            {
                "label": localize_template_text(item.get("label", "")),
                "path": localize_template_text(item.get("path", "")),
                "note": localize_template_text(item.get("note", "")),
            }
        )
    markdown_payload["assets"] = localized_assets
    markdown_payload["notes"] = [localize_template_text(item) for item in normalize_string_list(markdown_payload.get("notes"))]
    markdown_payload["sources"] = [localize_template_text(item) for item in normalize_string_list(markdown_payload.get("sources"))]

    for section in markdown_payload.get("sections", []) or []:
        section["heading"] = section.get("display_heading", section.get("heading", ""))
        section["instruction"] = localize_template_text(section.get("instruction", ""))
        section["paragraphs"] = [localize_template_text(item) for item in normalize_string_list(section.get("paragraphs"))]
        section["bullets"] = [localize_template_text(item) for item in normalize_string_list(section.get("bullets"))]
        section["numbered"] = [localize_template_text(item) for item in normalize_string_list(section.get("numbered"))]
        table = normalize_table(section.get("table"))
        table["title"] = localize_template_text(localized_table_title(table.get("title", "")))
        table["headers"] = [localize_template_text(localized_header_text(item)) for item in table.get("headers", [])]
        table["rows"] = [[localize_template_text(cell) for cell in row] for row in table.get("rows", [])]
        section["table"] = table
        evidence_refs = []
        for item in normalize_evidence_refs(section.get("evidence_refs")):
            evidence_refs.append(
                {
                    "source_type": localize_template_text(item.get("source_type", "")),
                    "source_id": localize_template_text(item.get("source_id", "")),
                    "source_url": localize_template_text(item.get("source_url", "")),
                    "time_range": localize_template_text(item.get("time_range", "")),
                    "excerpt": localize_template_text(item.get("excerpt", "")),
                    "supports": localize_template_text(item.get("supports", "")),
                }
            )
        section["evidence_refs"] = evidence_refs

    markdown = render_markdown_from_payload(markdown_payload)
    replacements = {
        "# Scene": "# 场景",
        "- Scene: ": "- 场景：",
        "- Project: ": "- 项目：",
        "- Deliverable Type: ": "- 交付物类型：",
        "- Generated: ": "- 生成时间：",
        "- Status: ": "- 状态：",
        "- Scenario File: ": "- 场景文件：",
        "## Working Context": "## 任务上下文",
        "### Inputs": "### 输入材料",
        "### Minimum Evidence": "### 最低证据要求",
        "### Ideal Evidence": "### 理想证据补充",
        "### Constraints": "### 约束条件",
        "### Requested Outputs": "### 目标交付",
        "### Ready Checklist": "### 开跑前检查",
        "## Executive Summary": "## 执行摘要",
        "- Conclusion: ": "- 核心结论：",
        "- Why It Matters: ": "- 为什么重要：",
        "- Next Action: ": "- 下一步动作：",
        "- Confidence: ": "- 置信度：",
        "## Operator Checklist": "## 操作检查清单",
        "## Common Failure Modes": "## 常见失败模式",
        "## Operator Guide": "## 操作指引",
        "## Direct-Use Template": "## 直接执行模板",
        "- Recommended Request: ": "- 推荐请求：",
        "- Recommended Request (ZH): ": "- 推荐请求（中文）：",
        "### Operator Dispatch": "### 操作分发",
        "- Runner Args:": "- 运行参数：",
        "### Variable Inputs": "### 可变输入",
        "| Variable | Meaning | Example | Required |": "| 变量 | 含义 | 示例 | 是否必填 |",
        "### Codex Prompt Scaffold": "### Codex 提示词骨架",
        "### Chinese Prompt Scaffold": "### 中文提示词骨架",
        "### Workflow Steps": "### 执行步骤",
        "### Output Checklist": "### 交付检查清单",
        "## Evidence": "## 证据总表",
        "| Label | Detail | Source |": "| 标签 | 详情 | 来源 |",
        "## Assets": "## 资产清单",
        "## Notes": "## 备注",
        "## Sources": "## 来源",
        "### Evidence References": "### 证据引用",
        "### 证据总表 References": "### 证据总表",
        "### Price-Band Differences": "### 价格带差异",
        "### Next Capture Upgrade": "### 下轮补采升级",
    "Use scene 01 as the governing workflow.": "以场景 01 作为本次工作的主流程。",
    "Use scene 08 as the governing workflow.": "以场景 08 作为本次工作的主流程。",
    "Use scene 18 as the governing workflow.": "以场景 18 作为本次工作的主流程。",
    "Use scene 19 as the governing workflow.": "以场景 19 作为本次工作的主流程。",
    "Download-ready video source": "可下载视频源",
    "The capture pack preserves playable or downloadable video detail for frame-by-frame review.": "该 capture pack 保留了可播放或可下载的视频细节，可继续用于逐帧复核。",
    "15 TikTok search results with links, metrics, and first-hook notes": "15 条 TikTok 搜索结果，含链接、基础指标与首拍钩子笔记",
    "ranked shortlist plus teardown priority": "排序短名单与深拆优先级",
    "Comments from 3 products with repeated phrases highlighted": "来自 3 个商品的评论样本，已标出重复短语与购买型语言",
    "pain-language synthesis and persona guidance": "买家语言提炼与人群画像启发",
    "2+ competitor accounts with one week's posts and performance notes": "2 个以上竞品账号的一周帖子与表现备注",
    "weekly competitor report and action board": "竞品周报与动作看板",
    "Recent post table with metrics, hooks, and content-type labels": "最近帖子表，包含指标、钩子和内容类型标签",
    "Recent post table with metrics, 钩子s, and content-type labels": "最近帖子表，包含指标、钩子和内容类型标签",
    "performance retro and next-cycle test plan": "表现复盘与下轮测试计划",
    "Separate category-level base value from category-level improvement opportunity.": "把品类级基础价值与品类级改进机会明确分开。",
        "| Source Type | Source ID | Source URL | Time Range | Excerpt | Supports |": "| 来源类型 | 来源标识 | 来源链接 | 时间范围 | 摘录 | 支撑结论 |",
        "_Add the user brief, market, product, and evidence notes here._": "_在此补充用户需求、市场、产品与证据备注。_",
        "_Fill this field._": "_待补充。_",
        "_Optional._": "_可选。_",
        "_Fill this section._": "_本章节待补充。_",
    }
    for source, target in replacements.items():
        markdown = markdown.replace(source, target)
    phrase_replacements = {
        "breakdown_report": "拆解报告",
        "insight_report": "洞察报告",
        "collection_board": "采集看板",
        "creation_brief": "创作制作简报",
        "testing_matrix": "测试矩阵",
        "imported": "已导入",
        "Human-readable run or campaign name": "便于识别的运行名或项目名",
        "Target market or locale when the scene depends on one market": "当场景依赖单一市场时的目标市场或地区",
        "Links, screenshots, transcripts, exports, OCR text, or copied notes used as source evidence": "作为源证据使用的链接、截图、转写、导出文件、OCR 文本或摘录备注",
        "What the operator wants this scene to produce": "操作者希望该场景产出的结果",
        "yes": "是",
        "recommended": "建议",
        "One video link or storyboard": "1 条视频链接或一份分镜摘要",
        "Transcript or subtitle notes": "转写稿或字幕笔记",
        "Frame notes or screenshots": "逐帧笔记或截图",
        "Optional basic performance context": "可选的基础表现数据",
        "Transcript or subtitle notes": "转写稿或字幕笔记",
        "Screenshots by beat": "按节拍截取的截图",
        "Basic performance context": "基础表现上下文",
        "Audio or BGM cue": "音频或 BGM 线索",
        "Timeline breakdown table": "时间轴拆解表",
        "Video-type classification": "视频类型归类",
        "BGM analysis": "BGM 分析",
        "Three-part viral interpretation": "三段式爆点解读",
        "Reusable mechanism": "可复用机制",
        "Adaptation advice": "改编建议",
        "Separate deep logic from surface style.": "把深层机制与表层风格分开看。",
        "If there is no voiceover, reconstruct the logic from subtitles, actions, cuts, and visual proof.": "如果没有口播，就从字幕、动作、剪辑与视觉证明链里重建逻辑。",
        "Hook, proof, and close can be reconstructed": "能重建 hook、证明与收口。",
        "At least one adaptation target is known": "至少已知 1 个改编目标。",
        "A timeline table can be filled without inventing missing beats": "可以在不臆造缺失节拍的前提下补完时间轴。",
        "Reconstruct the video in order with a time-ranged beat table.": "按顺序重建视频，并输出带时间范围的节拍表。",
        "Classify the video type before over-generalizing the lesson.": "在泛化结论前，先判断视频类型。",
        "Separate core mechanism from creator-specific surface style.": "把核心机制与创作者专属表层风格分开。",
        "Capture BGM, subtitle behavior, and transition rhythm explicitly.": "明确记录 BGM、字幕方式与转场节奏。",
        "Write at least one adaptation path before closing the report.": "在结束报告前至少写出 1 条改编路径。",
        "Confusing visual polish with the true conversion mechanism.": "把视觉精致感误当成真实转化机制。",
        "Ignoring no-voiceover logic because the transcript is sparse.": "因为转写稀疏就忽略无口播视频的成立逻辑。",
        "Skipping the close or CTA logic because it looks simple.": "因为收口或 CTA 看起来简单就跳过不拆。",
        "Giving abstract praise without reusable takeaways.": "只给抽象夸赞，没有可复用结论。",
        "The timeline reconstructs the video in order with evidence-backed beats.": "时间轴已按顺序重建，并有证据支撑各节拍。",
        "The hook, proof, and close are reconstructed in order.": "hook、证明与收口已按顺序重建。",
        "The 钩子, proof, and close are reconstructed in order.": "hook、证明与收口已按顺序重建。",
        "Hook, proof, and close can be reconstructed": "钩子、证明与收口可以被重建",
        "钩子, proof, and close can be reconstructed": "钩子、证明与收口可以被重建",
        "BGM, video type, and conversion rhythm are stated explicitly.": "BGM、视频类型与转化节奏已明确写出。",
        "The core mechanism is distinguished from surface style.": "已区分核心机制与表层风格。",
        "At least one adaptation path is concrete enough to produce from.": "至少有 1 条改编路径具体到可进入制作。",
        "Video evidence": "视频证据",
        "Link, screenshots, transcript notes, or manual reconstruction.": "链接、截图、转写笔记或人工重建内容。",
        "Audio evidence": "音频证据",
        "BGM name, audio style, or subtitle notes if no clean transcript exists.": "BGM 名称、音频风格，或无干净转写时的字幕笔记。",
        "Make a sharp judgment about why this single video works or fails.": "直接判断这条单视频为什么成立或为什么失效。",
        "Map the video from open to close with a beat-level reconstruction.": "按节拍重建视频从开头到结尾的完整结构。",
        "Describe the underlying mechanism, not just the visible style.": "描述底层机制，而不只写可见风格。",
        "Extract only the transferable parts.": "只提炼可迁移的部分。",
        "Interpret the viral logic using the strongest three practical lenses.": "用最实战的三种视角解释爆点逻辑。",
        "Describe how audio, BGM, subtitle density, and editing rhythm contribute to performance.": "说明音频、BGM、字幕密度与剪辑节奏如何影响表现。",
        "Turn the breakdown into a direct replication or editing blueprint.": "把拆解结果转成可直接复刻或剪辑执行的蓝图。",
        "Give one safer and one more aggressive adaptation path.": "分别给出一条更稳妥和一条更激进的改编路径。",
        "Reference video": "参考视频",
        "Screenshots or frame summary": "截图或逐帧摘要",
        "Transcript notes": "转写笔记",
        "Optional user product to adapt onto": "可选的用户产品改写目标",
        "Transcript": "转写稿",
        "Frame-by-frame notes": "逐帧笔记",
        "User product for adaptation": "用于改写的用户产品",
        "Scene or talent notes": "场景或人物备注",
        "One video or visual summary": "1 条视频或一份视觉摘要",
        "If evidence is thin, mark the prompt as low-confidence.": "如果证据偏薄，就把提示词标记为低置信度。",
        "Do not invent hidden production details; keep uncertain fields explicit.": "不要臆造隐藏制作细节，不确定字段要明确写出来。",
        "Inferred original brief": "反推原始制作简报",
        "Generator-ready schema": "可生成结构",
        "Shot-by-shot table": "分镜逐条表",
        "Product-adapted brief": "产品适配制作简报",
        "Field-level confidence flags": "字段级置信度标记",
        "Visual evidence is sufficient to infer shot language": "视觉证据足以反推镜头语言。",
        "Low-confidence gaps are explicit if evidence is thin": "如果证据偏薄，低置信度缺口已明确写出。",
        "The inferred brief can be separated into original and adapted versions": "反推制作简报已能拆成原版与改写版。",
        "Use scene 05 as the governing workflow.": "以场景 05 作为本次工作的主流程。",
        "State the likely creative intent before writing the inferred prompt.": "在写反推提示词前，先说明可能的创作意图。",
        "Translate observed output into generator-ready prompt blocks, not style buzzwords.": "把观察到的成品翻译成可直接喂给生成器的提示块，而不是空泛风格词。",
        "Keep an inferred-original brief separate from the product-adapted brief.": "把反推原始制作简报与产品改写制作简报分开写。",
        "Use field-level confidence labels where evidence is thin.": "在证据偏薄的字段上使用字段级置信度标签。",
        "Mark low-confidence guesses when evidence is thin.": "当证据偏薄时，要把猜测明确标成低置信度。",
        "Inventing prompt details not justified by the video.": "臆造视频里没有支撑的提示词细节。",
        "Only describing visual style without pacing, shot, and VO logic.": "只描述视觉风格，不写节奏、镜头和口播逻辑。",
        "Skipping shot-level structure and leaving only one generic prompt paragraph.": "跳过分镜级结构，只留下一个泛泛提示词段落。",
        "Forgetting to rewrite the inferred brief for the user's product.": "忘了按用户产品重写反推制作简报。",
        "The inferred brief is structured into generator-ready creation blocks.": "反推制作简报已按可生成创作模块结构化。",
        "Shot-level rows are present, not only one top-level prompt.": "已经给出分镜级条目，而不只是一个顶层提示词段落。",
        "An adapted version for the user's product is included when product context exists.": "在有产品上下文时，已经给出面向用户产品的改写版。",
        "Weakly supported inferences are clearly labeled.": "证据较弱的推断已清晰标注。",
        "The output can be adapted to a user product without redoing the analysis from zero.": "当前输出无需从零重做分析，就能继续改写到用户产品上。",
        "Visual evidence": "视觉证据",
        "Attach frames or describe the scene order.": "补上关键帧，或描述场景顺序。",
        "Audio / transcript evidence": "音频 / 转写证据",
        "Paste key spoken lines or subtitle notes.": "粘贴关键口播句或字幕笔记。",
        "Summarize the likely creative intent behind the piece.": "概括这条内容背后的可能创作意图。",
        "Reconstruct the generator-ready brief from observed output.": "从观察到的成品反推可直接生成的制作简报结构。",
        "State what makes the reconstructed brief effective and what is still inferred.": "说明这份反推制作简报为什么有效，以及哪些部分仍属推断。",
        "Write the inferred original prompt or creation brief in a generator-ready schema.": "把反推原始提示词 / 创作制作简报写成可直接生成的结构。",
        "Describe the shot-by-shot structure and where the inference is weak.": "描述分镜结构，并标出哪些推断仍偏弱。",
        "If a user product exists, state how to rewrite the brief for it.": "如果已有用户产品，请说明如何把制作简报改写过去。",
        "Make the brief directly usable by a generator or editor without another analysis round.": "让这份制作简报无需再二次分析，就能直接给生成器或剪辑执行。",
        "One creator account or several videos from one creator": "1 个创作者账号或同一创作者的多条视频",
        "Top videos": "高表现视频",
        "Transcripts": "转写稿",
        "Performance notes": "表现备注",
        "Optional posting-time and hashtag notes": "可选的发布时间与 hashtag 备注",
        "Posting-time or hashtag notes": "发布时间或 hashtag 备注",
        "Separate creator-specific advantage from transferable pattern.": "把创作者专属优势与可迁移模式分开。",
        "Account overview": "账号概览",
        "High vs low interaction comparison": "高低互动对比",
        "Repeatable formulas": "可重复公式",
        "Non-transferable advantages": "不可迁移优势",
        "New-script bridge": "新脚本桥接",
        "Repeated patterns appear across multiple videos": "重复模式已经跨多条视频出现。",
        "High- and low-interaction examples can be compared": "高低互动样本已经可以对比。",
        "Creator-specific advantages are separated": "创作者专属优势已经分开标注。",
        "Adaptation path for user product is possible": "已经具备向用户产品迁移的改写路径。",
        "Use scene 17 as the governing workflow.": "以场景 17 作为本次工作的主流程。",
        "Use multiple creator samples before declaring a repeatable formula.": "在宣称形成公式前，先使用多个创作者样本。",
        "Summarize the account baseline before extracting formulas.": "在提炼公式前，先总结账号基线。",
        "Compare high-performing and weak-performing samples directly.": "直接对比高表现与弱表现样本。",
        "Map repeated hook, pacing, proof, and CTA patterns separately.": "把重复出现的钩子、节奏、证明和 CTA 模式分开拆。",
        "Explicitly separate transferable pattern from creator advantage.": "明确区分可迁移模式与创作者优势。",
        "Overfitting one breakout video into a full creator formula.": "把单条爆款过拟合成完整创作者公式。",
        "Ignoring low-performing samples and therefore missing contrast.": "忽略低表现样本，导致缺少对比。",
        "Ignoring trust or identity advantages unique to the creator.": "忽略创作者独有的信任或身份优势。",
        "Ending with admiration instead of adaptation rules.": "最后停留在赞美，而不是落到改编规则。",
        "The account overview and content baseline are explicit before formula claims.": "在提出公式判断前，账号概览和内容基线已明确。",
        "Repeated patterns are supported by multiple creator samples.": "重复模式已经被多个创作者样本支撑。",
        "High- and low-interaction differences are compared directly.": "高低互动差异已直接对比。",
        "Hook and pacing formulas are reusable enough to seed new scripts.": "钩子与节奏公式已经足够可复用，可直接孵化新脚本。",
        "Transferable rules are separated from creator-only advantages.": "可迁移规则已与创作者独有优势分离。",
        "The report ends with adaptation guidance for a new product or account.": "报告最后已落到面向新产品或新账号的改写指引。",
        "Creator sample set": "创作者样本集",
        "List the creator's top or representative videos.": "列出创作者的高表现或代表性视频。",
        "Summarize the creator's repeatable winning pattern.": "概括创作者可重复的胜出模式。",
        "Map the creator's recurring content structure.": "梳理这位创作者反复出现的内容结构。",
        "Start with an account overview before extracting formulas.": "先做账号概览，再提炼公式。",
        "Compare high-interaction versus low-interaction content before declaring a formula.": "在宣布某个公式成立前，先比较高互动与低互动内容。",
        "Extract what can transfer to another account or product.": "提炼哪些部分可迁移到别的账号或产品。",
        "List the parts that depend on this specific creator and the reusable formulas worth copying.": "列出依赖该创作者本人的部分，以及值得复制的公式。",
        "Capture recurring non-script patterns that still affect performance.": "记录那些虽然不是脚本文案、但仍影响表现的重复模式。",
        "Describe how to migrate the pattern to the user's product or account.": "说明如何把这套模式迁移到用户自己的产品或账号。",
        "Comments from at least 2 products": "至少 2 个商品的评论",
        "Comments from 2+ products": "来自 2 个以上商品的评论",
        "Market": "市场",
        "Product positioning goal": "产品定位目标",
        "Optional price-band notes": "可选的价格带备注",
        "Purchase factor synthesis": "购买因素提炼",
        "Praise keyword synthesis": "好评关键词提炼",
        "Complaint pain-point synthesis": "差评痛点提炼",
        "Price-band difference view": "价格带差异视图",
        "Persona summary": "人群画像总结",
        "Selection and content implications": "选品与内容启发",
        "Real TikTok capture-pack import from ": "真实 TikTok capture-pack 导入自 ",
        " for uncategorized lane. Current board size: ": "，当前用于未分类赛道。当前看板规模：",
        " ranked / ": " 条已排序 / ",
        " qualified.": " 条达标。",
        "qualified with min-like threshold ": "条达标，最低点赞阈值 ",
        "Reference video frames, transcript snippets, and pacing notes": "参考视频逐帧画面、转写片段与节奏笔记",
        "Source profile: ": "来源账号：",
        "session quality: ": "会话质量：",
        "queries: ": "查询词：",
        "topics: ": "主题：",
        "If comment volume is light, mark findings as provisional.": "如果评论量偏少，应把结论标记为暂定判断。",
        "Conclusions should stay tied to ranked metrics, captions, and capture-pack summaries only.": "结论必须只绑定到排序指标、caption 和 capture-pack 摘要，不要外推。",
        "Comments stay grouped by product": "评论应按商品维度分组保留。",
        "Repeated phrases can be quoted": "可以直接引用重复出现的原话。",
        "Low-volume caveats are explicit": "样本量不足的前提已明确写出。",
        "Source product labels survive the merge": "合并后仍保留来源商品标记。",
        "Top-ranked videos are clearly identified": "已明确标出高排名视频。",
        "Transferable pattern is separated from profile-specific brand power": "已将可迁移模式与账号自身品牌势能区分。",
        "Keep comments grouped by product before merging category signals.": "合并品类信号前，先按商品维度保留评论分组。",
        "Preserve source-product labels all the way into the insight layer.": "在洞察层也要保留来源商品标签。",
        "Quote repeated user language, not only analyst paraphrases.": "优先保留重复出现的用户原话，而不是只给分析师转述。",
        "Translate pains and desires into product and script implications.": "把痛点和欲望翻译成产品与脚本层的启发。",
        "Mixing one-off complaints with true repeated pains.": "把一次性抱怨和真正重复痛点混在一起。",
        "Summarizing sentiment without concrete user phrases.": "只总结情绪，不保留具体用户话术。",
        "Collapsing source-product differences too early and losing price-band insight.": "过早抹平来源商品差异，丢失价格带洞察。",
        "Ignoring the difference between desire, complaint, and trust signal.": "忽略欲望、抱怨和信任信号之间的区别。",
        "Repeated pain, desire, and trust signals are separated clearly.": "已清晰区分重复的痛点、欲望与信任信号。",
        "Real user-language evidence is preserved.": "已保留真实用户语言证据。",
        "Source products remain visible through the merged analysis.": "合并分析后仍能看见来源商品。",
        "Persona and messaging implications follow directly from the mined comments.": "人群画像与话术启发均直接来自评论挖掘。",
    "Profile:": "账号：",
    "Ranked video count: ": "排序视频数：",
    "Qualified video count: ": "达标视频数：",
    "Capture root: ": "采集根目录：",
    "排序视频 count: ": "排序视频数：",
    "summary.json or aggregate_summary.json": "汇总文件（summary.json 或 aggregate_summary.json）",
    "profile_summary.json or summary.json": "账号汇总文件（profile_summary.json 或 summary.json）",
    "ranked_videos.json or aggregate_ranked_videos.json": "排序视频文件（ranked_videos.json 或 aggregate_ranked_videos.json）",
    "aggregate_qualified_videos.json or qualified_video_links.txt": "达标视频文件（aggregate_qualified_videos.json 或 qualified_video_links.txt）",
    "aggregate_report.md": "聚合报告（aggregate_report.md）",
    "video_details.json": "视频明细（video_details.json）",
    "1 条视频链接或一份分镜摘要 summary": "1 条视频链接或一份分镜摘要",
    "市场 context": "市场上下文",
    "排序方式 and shop-cart filter state": "排序方式与购物车筛选状态",
    "结论必须只绑定到排序指标、caption 和 capture-pack 摘要，不要外推。": "结论必须只绑定到排序指标、标题/钩子文本和采集包摘要，不要外推。",
    "真实 TikTok capture-pack 导入自 ": "真实 TikTok 采集包导入自 ",
    "capture-pack": "采集包",
    "capture pack": "采集包",
    "会话质量：unknown": "会话质量：待补",
    "会话质量：browser_same_origin_api_ok": "会话质量：浏览器同源接口正常",
    "会话质量：tikmatrix_profile_posts_export": "会话质量：TikMatrix 主页帖子导出",
    "查询词：none": "查询词：未提供",
    "主题：none": "主题：未提供",
    "会话=browser_same_origin_api_ok": "会话=浏览器同源接口正常",
    "会话=tikmatrix_profile_posts_export": "会话=TikMatrix 主页帖子导出",
    " likes, ": " 点赞, ",
    " plays, ": " 播放, ",
    " shares, ": " 分享, ",
    " comments.": " 评论.",
    "Topic text missing": "主题文本缺失",
    "Hook text missing": "钩子文本缺失",
    "not_detected": "未检测到",
    "sing your ️ out": "sing your heart out",
    "little moments ️": "little moments",
    "proud of you ️": "proud of you",
    "TikTok Validation Scene 05 Capture": "场景 05 校验样例",
    "| Summary |": "| 汇总 |",
    "| Profile summary |": "| 账号汇总 |",
        "Reference video frames, transcript snippets, and pacing notes": "参考视频逐帧画面、转写片段与节奏笔记",
        "参考视频 frames, transcript snippets, and pacing notes": "参考视频逐帧画面、转写片段与节奏笔记",
        "Use scene 04 as the governing workflow.": "以场景 04 作为本次工作的主流程。",
        "Use scene 08 as the governing workflow.": "以场景 08 作为本次工作的主流程。",
        "Normalize the provided evidence into this input set before analysis: ": "分析前先把现有证据归整为以下输入：",
        "If evidence is missing, state the gap explicitly before continuing. Minimum evidence to proceed: ": "如果证据不足，先明确缺口再继续。最低开工证据：",
        "Produce these outputs in operator-ready form: ": "最终必须产出以下可直接给运营使用的结果：",
        "Fill the scaffold with reusable conclusions, tables, ranking logic, and next actions instead of generic commentary.": "优先填入可复用结论、表格、排序逻辑和下一步动作，不要停留在泛泛点评。",
        "Summarize what these comments reveal about the category buyer.": "概括这些评论揭示出的品类购买者特征。",
        "State the strongest demand-side insight.": "写出最强的需求侧判断。",
        "Cluster repeated user language across products.": "跨商品聚类重复出现的用户语言。",
        "Turn the user language into next decisions.": "把用户语言转成下一步决策。",
        "List missing evidence or weak conclusions.": "列出缺失证据或薄弱结论。",
    }
    for source, target in phrase_replacements.items():
        markdown = markdown.replace(source, target)
    return markdown


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a structured TikTok Growth Operator scene report JSON into Markdown, DOCX, and XLSX outputs."
    )
    parser.add_argument("--input", help="Structured report JSON path.")
    parser.add_argument("--scene", help="Scene id or slug for scaffold generation.")
    parser.add_argument("--project", help="Project name for scaffold generation.")
    parser.add_argument("--context-file", help="Optional UTF-8 context file for scaffold generation.")
    parser.add_argument("--output-dir", required=True, help="Directory for rendered outputs.")
    parser.add_argument(
        "--formats",
        default="md,docx,xlsx",
        help="Comma-separated output formats: md, docx, xlsx.",
    )
    parser.add_argument(
        "--base-name",
        default="",
        help="Optional explicit output base filename without extension.",
    )
    return parser.parse_args()


def load_json(path: Path) -> dict:
    loaded = read_json_file(path)
    if not isinstance(loaded, dict):
        raise SystemExit(f"Structured report JSON must be an object: {path}")
    return loaded


def slugify(text: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9_-]+", "-", normalize_text(text).lower())
    return normalized.strip("-") or "report"


def normalize_string_list(values: list | None) -> list[str]:
    return [normalize_text(item) for item in values or [] if normalize_text(item)]


def is_local_path_text(value: str) -> bool:
    text = normalize_text(value or "")
    return bool(text) and (
        text.startswith("\\\\")
        or (len(text) > 2 and text[1] == ":" and text[2] in {"\\", "/"})
    )


def compact_display_path(value: str) -> str:
    text = normalize_text(value or "")
    if not is_local_path_text(text):
        return text
    try:
        path = Path(text)
        parts = path.parts
        if len(parts) >= 3:
            return Path(*parts[-3:]).as_posix()
        return path.as_posix()
    except (OSError, ValueError):
        return text


EMBEDDED_WINDOWS_PATH_RE = re.compile(r"(?<![A-Za-z])(?P<path>[A-Za-z]:[\\/][^|;\n\r]+)")


def compact_embedded_paths(value: str) -> str:
    text = normalize_text(value or "")
    if not text:
        return ""

    def _replace(match: re.Match[str]) -> str:
        candidate = match.group("path").strip().rstrip(".,)")
        compact = compact_display_path(candidate)
        suffix = match.group("path")[len(candidate):]
        return compact + suffix

    return EMBEDDED_WINDOWS_PATH_RE.sub(_replace, text)


def likely_dirty_zh_text(value: str) -> bool:
    text = normalize_text(value or "")
    if not text:
        return False
    if not any("\u4e00" <= ch <= "\u9fff" for ch in text):
        return False
    noise_markers = ("鎴", "锛", "銆", "馃", "", "杈", "缁", "浜", "鍏", "鏈")
    return any(marker in text for marker in noise_markers)


def display_clean_text(value: str) -> str:
    text = normalize_text(value or "")
    if not text:
        return ""
    if is_local_path_text(text):
        return compact_display_path(text)
    return compact_embedded_paths(text)


def normalize_execution_template_display(execution_template: dict) -> dict:
    recommended_request_zh = localize_template_text(normalize_text(execution_template.get("recommended_request_zh", "")))
    prompt_scaffold_zh = [localize_template_text(item) for item in normalize_string_list(execution_template.get("codex_prompt_scaffold_zh"))]
    if likely_dirty_zh_text(recommended_request_zh):
        recommended_request_zh = ""
    if prompt_scaffold_zh and any(likely_dirty_zh_text(item) for item in prompt_scaffold_zh):
        prompt_scaffold_zh = []
    return {
        "recommended_request": localize_template_text(execution_template.get("recommended_request", "")),
        "recommended_request_zh": recommended_request_zh,
        "recommended_runner_args": normalize_string_list(execution_template.get("recommended_runner_args")),
        "variable_inputs": execution_template.get("variable_inputs", []),
        "codex_prompt_scaffold": [localize_template_text(item) for item in normalize_string_list(execution_template.get("codex_prompt_scaffold"))],
        "codex_prompt_scaffold_zh": prompt_scaffold_zh,
        "workflow_steps": [localize_template_text(item) for item in normalize_string_list(execution_template.get("workflow_steps"))],
        "output_checklist": [localize_template_text(item) for item in normalize_string_list(execution_template.get("output_checklist"))],
    }


def normalize_table(table: dict | None) -> dict:
    payload = table or {}
    headers = [localized_header_text(display_clean_text(item)) for item in normalize_string_list(payload.get("headers"))]
    rows = []
    for row in payload.get("rows", []) or []:
        normalized_row = []
        for cell in row:
            cell_text = display_clean_text(normalize_text(cell))
            normalized_row.append(localize_template_text(cell_text))
        rows.append(normalized_row)
    return {
        "title": localize_template_text(localized_table_title(display_clean_text(normalize_text(payload.get("title", ""))))),
        "headers": headers,
        "rows": rows,
    }


def normalize_evidence_refs(rows: list[dict] | None) -> list[dict]:
    normalized_rows = []
    for item in rows or []:
        normalized_rows.append(
            {
                "source_type": localize_template_text(display_clean_text(str(item.get("source_type", "")).strip())),
                "source_id": localize_template_text(display_clean_text(str(item.get("source_id", "")).strip())),
                "source_url": localize_template_text(display_clean_text(str(item.get("source_url", "")).strip())),
                "time_range": localize_template_text(display_clean_text(str(item.get("time_range", "")).strip())),
                "excerpt": localize_template_text(display_clean_text(str(item.get("excerpt", "")).strip())),
                "supports": localize_template_text(display_clean_text(str(item.get("supports", "")).strip())),
            }
        )
    return normalized_rows


def resolve_payload(args: argparse.Namespace) -> dict:
    if args.input:
        payload = load_json(Path(args.input))
    else:
        if not args.scene or not args.project:
            raise SystemExit("Provide --input, or provide both --scene and --project.")
        skill_root = Path(__file__).resolve().parents[1]
        catalog = load_catalog(skill_root)
        scene = resolve_scene(catalog, args.scene)
        context = read_utf8_text(Path(args.context_file)) if args.context_file else ""
        payload = build_report_payload(scene, args.project, context)
    payload = normalize_nested(payload)

    metadata = payload.setdefault("metadata", {})
    scene_id = normalize_text(metadata.get("scene", ""))
    skill_root = Path(__file__).resolve().parents[1]
    catalog = load_catalog(skill_root)
    if scene_id:
        scene = resolve_scene(catalog, scene_id)
        metadata.setdefault("scene_slug", scene["slug"])
        metadata.setdefault("scene_title", scene["title"])
        metadata.setdefault("deliverable_type", scene["deliverable_type"])
        metadata.setdefault("scenario_file", scene["scenario_file"])
    metadata.setdefault("project", "untitled-project")
    metadata.setdefault("title", f"Scene {metadata.get('scene', 'XX')} Report - {metadata['project']}")
    metadata.setdefault("status", "draft")
    metadata.setdefault("generated_at", "")

    working_context = payload.setdefault("working_context", {})
    working_context["summary"] = localize_template_text(working_context.get("summary", ""))
    working_context["inputs"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("inputs"))]
    working_context["minimum_evidence"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("minimum_evidence"))]
    working_context["ideal_evidence"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("ideal_evidence"))]
    working_context["constraints"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("constraints"))]
    working_context["requested_outputs"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("requested_outputs"))]
    working_context["ready_checklist"] = [localize_template_text(item) for item in normalize_string_list(working_context.get("ready_checklist"))]

    executive = payload.setdefault("executive_summary", {})
    executive["conclusion"] = localize_template_text(executive.get("conclusion", ""))
    executive["why_it_matters"] = localize_template_text(executive.get("why_it_matters", ""))
    executive["next_action"] = localize_template_text(executive.get("next_action", ""))
    executive["confidence"] = localize_template_text(executive.get("confidence", ""))

    operator_guide = payload.setdefault("operator_guide", {})
    operator_guide.setdefault("operator_checklist", [])
    operator_guide.setdefault("common_failure_modes", [])
    operator_guide["operator_checklist"] = [
        localize_template_text(item) for item in normalize_string_list(operator_guide.get("operator_checklist"))
    ]
    operator_guide["common_failure_modes"] = [
        localize_template_text(item) for item in normalize_string_list(operator_guide.get("common_failure_modes"))
    ]

    execution_template = payload.setdefault("execution_template", {})
    execution_template.setdefault("recommended_request", "")
    execution_template.setdefault("recommended_request_zh", "")
    execution_template.setdefault("recommended_runner_args", [])
    execution_template.setdefault("variable_inputs", [])
    execution_template.setdefault("codex_prompt_scaffold", [])
    execution_template.setdefault("codex_prompt_scaffold_zh", [])
    execution_template.setdefault("workflow_steps", [])
    execution_template.setdefault("output_checklist", [])
    normalized_variable_inputs = []
    for item in execution_template.get("variable_inputs", []) or []:
        item_name = normalize_text(item.get("name", ""))
        example_text = normalize_text(item.get("example", ""))
        if item_name == "project_name":
            example_text = localize_project_text(example_text) or example_text
        normalized_variable_inputs.append(
            {
                "name": item_name,
                "meaning": localize_template_text(item.get("meaning", "")),
                "example": localize_template_text(example_text),
                "required": localize_template_text(item.get("required", "")),
            }
        )
    execution_template_display = normalize_execution_template_display(execution_template)
    execution_template_display["variable_inputs"] = normalized_variable_inputs
    payload["execution_template"] = execution_template_display

    normalized_sections = []
    for section in payload.get("sections", []) or []:
        normalized_sections.append(
            {
                "heading": str(section.get("heading", "")).strip() or "Untitled Section",
                "display_heading": localized_section_heading(str(section.get("heading", "")).strip() or "Untitled Section"),
                "instruction": localize_template_text(str(section.get("instruction", "")).strip()),
                "paragraphs": [localize_template_text(item) for item in normalize_string_list(section.get("paragraphs"))],
                "bullets": [localize_template_text(item) for item in normalize_string_list(section.get("bullets"))],
                "numbered": [localize_template_text(item) for item in normalize_string_list(section.get("numbered"))],
                "evidence_refs": normalize_evidence_refs(section.get("evidence_refs")),
                "table": normalize_table(section.get("table")),
            }
        )
    payload["sections"] = normalized_sections

    normalized_evidence = []
    for item in payload.get("evidence", []) or []:
        normalized_evidence.append(
            {
                "label": localize_template_text(display_clean_text(str(item.get("label", "")).strip())),
                "detail": localize_template_text(display_clean_text(str(item.get("detail", "")).strip())),
                "source": localize_template_text(display_clean_text(str(item.get("source", "")).strip())),
            }
        )
    payload["evidence"] = normalized_evidence

    normalized_assets = []
    for item in payload.get("assets", []) or []:
        normalized_assets.append(
            {
                "label": localize_template_text(display_clean_text(str(item.get("label", "")).strip())),
                "path": localize_template_text(display_clean_text(str(item.get("path", "")).strip())),
                "note": localize_template_text(display_clean_text(str(item.get("note", "")).strip())),
            }
        )
    payload["assets"] = normalized_assets
    payload["notes"] = [localize_template_text(display_clean_text(item)) for item in normalize_string_list(payload.get("notes"))]
    payload["sources"] = [localize_template_text(display_clean_text(item)) for item in normalize_string_list(payload.get("sources"))]
    return payload


def infer_base_name(report: dict, explicit_base_name: str) -> str:
    if explicit_base_name.strip():
        return explicit_base_name.strip()
    metadata = report["metadata"]
    return f"scene-{metadata.get('scene', 'xx')}-{slugify(metadata.get('project', '') or metadata.get('title', 'report'))}"


def build_output_path(output_dir: Path, base_name: str, suffix: str) -> Path:
    candidate = output_dir / f"{base_name}{suffix}"
    try:
        path_text = str(candidate.resolve(strict=False))
    except OSError:
        path_text = str(candidate)
    max_path_chars = 240
    if len(path_text) <= max_path_chars:
        return candidate
    reserve = len(suffix) + 1
    available = max(24, max_path_chars - len(str(output_dir)) - reserve)
    trimmed = base_name[:available].rstrip("-_. ")
    if not trimmed:
        trimmed = "scene-report"
    return output_dir / f"{trimmed}{suffix}"


def has_execution_template(report: dict) -> bool:
    execution_template = report.get("execution_template", {}) or {}
    return any(
        execution_template.get(key)
        for key in [
            "recommended_request",
            "recommended_request_zh",
            "recommended_runner_args",
            "variable_inputs",
            "codex_prompt_scaffold",
            "codex_prompt_scaffold_zh",
            "workflow_steps",
            "output_checklist",
        ]
    )


def summary_preview_lines(text: str, limit: int = 4) -> list[str]:
    preview = []
    for raw_line in normalize_text(text or "", strip=False).splitlines():
        line = raw_line.strip()
        if line:
            preview.append(line)
        if len(preview) >= limit:
            break
    return preview


def estimate_text_line_count(value: object, chars_per_line: int = 48) -> int:
    text = normalize_text(value or "", strip=False)
    if not text.strip():
        return 1
    total = 0
    for raw_line in text.splitlines() or [""]:
        line = raw_line or " "
        total += max(1, (len(line) + chars_per_line - 1) // chars_per_line)
    return max(total, 1)


def set_sheet_row_height(
    ws,
    row: int,
    values: list[object],
    chars_per_line: int = 48,
    min_height: float = 18,
    max_height: float = 90,
) -> None:
    max_lines = max((estimate_text_line_count(value, chars_per_line) for value in values), default=1)
    ws.row_dimensions[row].height = max(min_height, min(max_height, 15 * max_lines))


def infer_sheet_column_widths(
    headers: list[str],
    rows: list[list[str]],
    min_width: int = 14,
    max_width: int = 56,
    preferred_widths: list[float] | None = None,
    per_column_max: list[float] | None = None,
) -> list[float]:
    widths: list[float] = []
    for column_index, header in enumerate(headers):
        max_len = len(str(header))
        for row in rows[:200]:
            if column_index >= len(row):
                continue
            text = normalize_text(row[column_index], strip=False)
            longest_line = max((len(line) for line in text.splitlines()), default=0)
            max_len = max(max_len, min(longest_line, 80))
        computed = float(min(max_width, max(min_width, round(max_len * 1.05 + 2))))
        if per_column_max and column_index < len(per_column_max):
            computed = min(computed, float(per_column_max[column_index]))
        if preferred_widths and column_index < len(preferred_widths):
            computed = max(computed, float(preferred_widths[column_index]))
        widths.append(computed)
    return widths


def list_sheet_layout(title: str) -> dict[str, object]:
    layout_map: dict[str, dict[str, object]] = {
        "Evidence": {
            "preferred_widths": [24.0, 46.0, 34.0],
            "per_column_max": [30.0, 52.0, 40.0],
            "chars_per_line": 24,
        },
        "Assets": {
            "preferred_widths": [22.0, 46.0, 34.0],
            "per_column_max": [28.0, 52.0, 38.0],
            "chars_per_line": 24,
        },
        "Notes": {
            "preferred_widths": [72.0],
            "per_column_max": [84.0],
            "chars_per_line": 44,
        },
        "Sources": {
            "preferred_widths": [72.0],
            "per_column_max": [84.0],
            "chars_per_line": 44,
        },
    }
    return layout_map.get(
        title,
        {
            "preferred_widths": [],
            "per_column_max": [],
            "chars_per_line": 26,
        },
    )


def infer_doc_table_widths(headers: list[str], rows: list[list[str]], usable_width: float = 6.8) -> list[float]:
    weights: list[float] = []
    for column_index, header in enumerate(headers):
        max_len = len(header)
        for row in rows[:40]:
            if column_index >= len(row):
                continue
            text = normalize_text(row[column_index], strip=False)
            longest_line = max((len(line) for line in text.splitlines()), default=0)
            max_len = max(max_len, min(longest_line, 60))
        weights.append(max(8.0, min(float(max_len), 42.0)))
    total = sum(weights) or 1.0
    widths = [round(max(0.75, usable_width * (weight / total)), 2) for weight in weights]
    if len(widths) > 1:
        widest_index = max(range(len(widths)), key=lambda index: widths[index])
        diff = round(usable_width - sum(widths), 2)
        widths[widest_index] = round(max(0.75, widths[widest_index] + diff), 2)
    return widths


CREATIVE_BRIEF_SCENE_IDS = frozenset({"09", "10", "13", "15", "16"})
CREATIVE_MATRIX_SCENE_IDS = frozenset({"11", "12", "14"})


def scene_doc_visual_theme(scene_id: str) -> dict[str, str]:
    if scene_id in CREATIVE_BRIEF_SCENE_IDS:
        return {
            "kicker": "创意制作项目卡",
            "accent": "6B4E71",
            "card_accent": "4A2C4F",
            "title_fill": "6B4E71",
        }
    if scene_id in CREATIVE_MATRIX_SCENE_IDS:
        return {
            "kicker": "创意流程项目卡",
            "accent": "2E6F5E",
            "card_accent": "1F4D42",
            "title_fill": "2E6F5E",
        }
    return {
        "kicker": "TikTok 增长运营项目卡",
        "accent": "355C7D",
        "card_accent": "1F4E78",
        "title_fill": "1F4E78",
    }


def creative_scene_section_layout_map() -> dict[tuple[str, str], dict[str, object]]:
    message_layout = {
        "doc_widths": [1.0, 1.55, 1.55, 1.7],
        "xlsx_preferred_widths": [14.0, 24.0, 24.0, 28.0],
        "xlsx_per_column_max": [16.0, 30.0, 30.0, 34.0],
        "xlsx_chars_per_line": 22,
    }
    structure_layout = {
        "doc_widths": [0.65, 1.35, 0.95, 1.35, 1.25, 1.15],
        "xlsx_preferred_widths": [10.0, 24.0, 16.0, 24.0, 22.0, 20.0],
        "xlsx_per_column_max": [12.0, 30.0, 18.0, 30.0, 26.0, 24.0],
        "xlsx_chars_per_line": 20,
    }
    constraints_layout = {
        "doc_widths": [1.15, 0.95, 1.55, 1.35],
        "xlsx_preferred_widths": [18.0, 16.0, 26.0, 22.0],
        "xlsx_per_column_max": [20.0, 18.0, 30.0, 26.0],
        "xlsx_chars_per_line": 22,
    }
    production_handoff_layout = {
        "doc_widths": [1.25, 1.65, 1.15, 1.15],
        "xlsx_preferred_widths": [20.0, 28.0, 18.0, 18.0],
        "xlsx_per_column_max": [22.0, 32.0, 20.0, 20.0],
        "xlsx_chars_per_line": 22,
    }
    variable_matrix_layout = {
        "doc_widths": [0.95, 1.15, 1.35, 1.1, 0.85, 1.15, 0.95],
        "xlsx_preferred_widths": [16.0, 18.0, 24.0, 18.0, 14.0, 18.0, 16.0],
        "xlsx_per_column_max": [18.0, 20.0, 30.0, 20.0, 16.0, 22.0, 18.0],
        "xlsx_chars_per_line": 20,
    }
    what_to_learn_layout = {
        "doc_widths": [1.2, 1.15, 1.0, 1.35, 1.0, 1.0],
        "xlsx_preferred_widths": [18.0, 18.0, 16.0, 24.0, 18.0, 18.0],
        "xlsx_per_column_max": [20.0, 20.0, 18.0, 28.0, 20.0, 20.0],
        "xlsx_chars_per_line": 22,
    }
    execution_handoff_layout = {
        "doc_widths": [1.35, 1.0, 1.0, 1.45],
        "xlsx_preferred_widths": [22.0, 16.0, 16.0, 26.0],
        "xlsx_per_column_max": [26.0, 18.0, 18.0, 30.0],
        "xlsx_chars_per_line": 22,
    }
    entries: dict[tuple[str, str], dict[str, object]] = {}
    for scene_id in CREATIVE_BRIEF_SCENE_IDS:
        entries[(scene_id, "Message")] = message_layout
        entries[(scene_id, "Structure")] = structure_layout
        entries[(scene_id, "Creative Constraints")] = constraints_layout
        entries[(scene_id, "Production Handoff")] = production_handoff_layout
    for scene_id in CREATIVE_MATRIX_SCENE_IDS:
        entries[(scene_id, "Variable Matrix")] = variable_matrix_layout
        entries[(scene_id, "What To Learn")] = what_to_learn_layout
        entries[(scene_id, "Execution Handoff")] = execution_handoff_layout
        if scene_id == "14":
            entries[(scene_id, "Production Handoff")] = production_handoff_layout
    return entries


def first_table_cell(section: dict | None, row_index: int = 0, column_index: int = 0) -> str:
    if not section:
        return ""
    rows = ((section.get("table") or {}).get("rows") or [])
    if row_index >= len(rows):
        return ""
    row = rows[row_index]
    if column_index >= len(row):
        return ""
    return normalize_text(row[column_index])


def scene_section_layout(report: dict, section: dict) -> dict[str, object]:
    scene_id = normalize_text(report.get("metadata", {}).get("scene", ""))
    heading = normalize_text(section.get("heading", ""))
    creative_layout = creative_scene_section_layout_map().get((scene_id, heading))
    if creative_layout is not None:
        return creative_layout
    layout_map: dict[tuple[str, str], dict[str, object]] = {
        ("01", "Objects To Track"): {
            "doc_widths": [0.7, 1.55, 1.65, 1.1, 1.15, 1.15, 0.9, 0.85, 0.75],
            "xlsx_preferred_widths": [10.0, 22.0, 22.0, 18.0, 22.0, 22.0, 18.0, 18.0, 12.0],
            "xlsx_per_column_max": [12.0, 28.0, 28.0, 20.0, 28.0, 28.0, 20.0, 22.0, 14.0],
            "xlsx_chars_per_line": 20,
        },
        ("01", "Why They Matter"): {
            "doc_widths": [1.45, 1.65, 1.35, 1.55, 1.8, 1.4],
            "xlsx_preferred_widths": [22.0, 22.0, 18.0, 20.0, 28.0, 22.0],
            "xlsx_per_column_max": [26.0, 26.0, 20.0, 24.0, 34.0, 26.0],
            "xlsx_chars_per_line": 20,
        },
        ("08", "High-Level Judgment"): {
            "doc_widths": [1.55, 1.0, 0.75, 1.85, 1.65],
            "xlsx_preferred_widths": [24.0, 14.0, 10.0, 28.0, 30.0],
            "xlsx_per_column_max": [28.0, 16.0, 12.0, 32.0, 34.0],
            "xlsx_chars_per_line": 20,
        },
        ("08", "Evidence Clusters"): {
            "doc_widths": [1.05, 2.35, 1.15, 1.45, 2.0],
            "xlsx_preferred_widths": [18.0, 38.0, 18.0, 24.0, 34.0],
            "xlsx_per_column_max": [20.0, 42.0, 20.0, 28.0, 38.0],
            "xlsx_chars_per_line": 22,
        },
        ("08", "Recommended Action"): {
            "doc_widths": [1.2, 2.3, 2.0, 1.3],
            "xlsx_preferred_widths": [18.0, 34.0, 28.0, 18.0],
            "xlsx_per_column_max": [20.0, 38.0, 34.0, 20.0],
            "xlsx_chars_per_line": 22,
        },
        ("08", "Open Questions"): {
            "doc_widths": [1.2, 1.9, 1.9, 2.0],
            "xlsx_preferred_widths": [18.0, 24.0, 24.0, 30.0],
            "xlsx_per_column_max": [20.0, 28.0, 28.0, 34.0],
            "xlsx_chars_per_line": 22,
        },
        ("18", "Objects To Track"): {
            "doc_widths": [1.05, 0.7, 1.55, 1.2, 0.95, 0.75, 0.95],
            "xlsx_preferred_widths": [18.0, 12.0, 28.0, 18.0, 18.0, 14.0, 18.0],
            "xlsx_per_column_max": [20.0, 14.0, 32.0, 20.0, 20.0, 16.0, 20.0],
            "xlsx_chars_per_line": 20,
        },
        ("18", "Why They Matter"): {
            "doc_widths": [1.2, 1.0, 1.8, 1.65, 1.75],
            "xlsx_preferred_widths": [18.0, 18.0, 30.0, 26.0, 28.0],
            "xlsx_per_column_max": [20.0, 20.0, 34.0, 30.0, 32.0],
            "xlsx_chars_per_line": 22,
        },
        ("18", "Fields To Capture Next Time"): {
            "doc_widths": [1.45, 2.35, 1.1],
            "xlsx_preferred_widths": [22.0, 38.0, 12.0],
            "xlsx_per_column_max": [24.0, 42.0, 14.0],
            "xlsx_chars_per_line": 24,
        },
        ("18", "Next Action"): {
            "doc_widths": [1.0, 2.5, 0.8, 2.5],
            "xlsx_preferred_widths": [14.0, 34.0, 10.0, 34.0],
            "xlsx_per_column_max": [16.0, 38.0, 12.0, 38.0],
            "xlsx_chars_per_line": 22,
        },
        ("19", "High-Level Judgment"): {
            "doc_widths": [1.15, 1.65, 2.2, 1.8],
            "xlsx_preferred_widths": [16.0, 24.0, 34.0, 28.0],
            "xlsx_per_column_max": [18.0, 28.0, 38.0, 32.0],
            "xlsx_chars_per_line": 22,
        },
        ("19", "Evidence Clusters"): {
            "doc_widths": [1.15, 1.5, 2.2, 1.8, 1.45],
            "xlsx_preferred_widths": [16.0, 22.0, 34.0, 28.0, 20.0],
            "xlsx_per_column_max": [18.0, 26.0, 40.0, 32.0, 22.0],
            "xlsx_chars_per_line": 22,
        },
        ("19", "Recommended Action"): {
            "doc_widths": [1.0, 2.2, 2.2, 1.4],
            "xlsx_preferred_widths": [14.0, 30.0, 30.0, 18.0],
            "xlsx_per_column_max": [16.0, 34.0, 34.0, 20.0],
            "xlsx_chars_per_line": 22,
        },
        ("19", "Open Questions"): {
            "doc_widths": [1.45, 2.1, 1.7, 1.55],
            "xlsx_preferred_widths": [22.0, 30.0, 22.0, 24.0],
            "xlsx_per_column_max": [24.0, 34.0, 26.0, 28.0],
            "xlsx_chars_per_line": 22,
        },
    }
    return layout_map.get(
        (scene_id, heading),
        {
            "doc_widths": None,
            "xlsx_preferred_widths": [],
            "xlsx_per_column_max": [],
            "xlsx_chars_per_line": 24,
        },
    )


def set_doc_font(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def shade_doc_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn_shared("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_doc_cell_vertical_alignment(cell, value: str = "center") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn_shared("w:val"), value)
    tc_pr.append(v_align)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def style_doc_table(table, header_row: bool = True, label_col: int | None = None) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_doc_cell_vertical_alignment(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
            if row_index == 0 and header_row:
                shade_doc_cell(cell, "D9EAF7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            elif label_col is not None and col_index == label_col:
                shade_doc_cell(cell, "EEF4FB")
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def style_doc_card_table(table, accent_fill: str = "D9EAF7") -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_doc_cell_vertical_alignment(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
            if row_index == 0:
                shade_doc_cell(cell, accent_fill)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            elif col_index == 0:
                shade_doc_cell(cell, "EEF4FB")
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def add_doc_kicker(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)


def add_doc_section_divider(document: Document, title: str, subtitle: str = "") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    if subtitle:
        sub = document.add_paragraph()
        sub.paragraph_format.space_after = Pt(8)
        sub_run = sub.add_run(subtitle)
        sub_run.italic = True
        sub_run.font.size = Pt(9.5)


def add_doc_badge_row(document: Document, badges: list[tuple[str, str]]) -> None:
    if not badges:
        return
    table = document.add_table(rows=2, cols=len(badges))
    for index, (label, value) in enumerate(badges):
        table.cell(0, index).text = label
        table.cell(1, index).text = value
    style_doc_card_table(table, accent_fill="1F4E78")
    widths = [max(1.1, round(6.6 / max(len(badges), 1), 2))] * len(badges)
    set_doc_table_widths(table, widths)


def add_doc_metric_grid(document: Document, metrics: list[tuple[str, str]], cols: int = 3) -> None:
    if not metrics:
        return
    rows = (len(metrics) + cols - 1) // cols
    table = document.add_table(rows=rows * 2, cols=cols)
    index = 0
    for row_group in range(rows):
        label_row = row_group * 2
        value_row = label_row + 1
        for col in range(cols):
            if index >= len(metrics):
                break
            label, value = metrics[index]
            table.cell(label_row, col).text = label
            table.cell(value_row, col).text = value
            index += 1
    style_doc_card_table(table, accent_fill="355C7D")
    set_doc_table_widths(table, [2.2] * cols)


def add_doc_spotlight_table(
    document: Document,
    title: str,
    subtitle: str,
    rows: list[tuple[str, str]],
    *,
    accent_fill: str = "1F4E78",
    widths: list[float] | None = None,
) -> None:
    add_doc_section_divider(document, title, subtitle)
    table = document.add_table(rows=max(1, len(rows)), cols=2)
    normalized_rows = rows or [("状态", "待补")]
    for index, (label, value) in enumerate(normalized_rows):
        table.cell(index, 0).text = label
        table.cell(index, 1).text = value or "待补"
    style_doc_table(table, header_row=False, label_col=0)
    if normalized_rows:
        shade_doc_cell(table.cell(0, 0), accent_fill)
        for run in table.cell(0, 0).paragraphs[0].runs:
            run.font.color.rgb = None
    set_doc_table_widths(table, widths or [1.45, 5.35])


def scene_cover_spotlight_rows(report: dict) -> list[tuple[str, str]]:
    scene_id = normalize_text(report.get("metadata", {}).get("scene", ""))
    executive = report.get("executive_summary", {})
    sections = {normalize_text(section.get("heading", "")): section for section in report.get("sections", [])}
    if scene_id == "18":
        objects = ((sections.get("Objects To Track") or {}).get("table") or {}).get("rows") or []
        actions = ((sections.get("Next Action") or {}).get("table") or {}).get("rows") or []
        strongest = objects[0][4] if objects and len(objects[0]) > 4 else executive.get("conclusion", "")
        compare_window = ""
        for paragraph in (sections.get("Executive Conclusion") or {}).get("paragraphs", []):
            if "比较窗口" in normalize_text(paragraph):
                compare_window = paragraph
                break
        action_line = "；".join([f"{row[0]}：{row[1]}" for row in actions[:3] if len(row) >= 2]) or (executive.get("next_action", "").strip())
        return [
            ("周报模式", "多账号竞品周报卡" if objects and len(objects[0]) >= 7 else "账号周基线卡"),
            ("比较窗口", compare_window or "比较窗口待补"),
            ("本周最强包装线", strongest or "本周最强包装线待补"),
            ("运营动作分发", action_line or "继续追 / 借鉴 / 减少跟进 待补"),
        ]
    if scene_id == "19":
        judgments = ((sections.get("High-Level Judgment") or {}).get("table") or {}).get("rows") or []
        tests = ((sections.get("Open Questions") or {}).get("table") or {}).get("rows") or []
        best_window = ""
        more_less = ""
        for row in judgments:
            if row and "最佳发布时间窗" in normalize_text(row[0]):
                best_window = row[1] if len(row) > 1 else ""
            if row and "多做" in normalize_text(row[0]):
                more_less = "；".join([cell for cell in row[:3] if normalize_text(cell)])
        test_line = "；".join([f"{row[0]}：{row[2]}" for row in tests[:2] if len(row) >= 3]) or "下轮测试计划待补"
        return [
            ("复盘模式", "账号复盘优化卡"),
            ("最佳发布时间窗", best_window or "最佳发布时间窗待补"),
            ("多做 / 少做", more_less or executive.get("next_action", "").strip() or "动作结论待补"),
            ("下轮测试", test_line),
        ]
    if scene_id in CREATIVE_BRIEF_SCENE_IDS:
        message = sections.get("Message") or {}
        structure = sections.get("Structure") or {}
        handoff = sections.get("Production Handoff") or {}
        hook_line = first_table_cell(message, 0, 1) or first_table_cell(message, 0, 0)
        proof_line = first_table_cell(structure, 1, 1) or first_table_cell(structure, 0, 1)
        owner_line = first_table_cell(handoff, 0, 2) or first_table_cell(handoff, 0, 0)
        return [
            ("创意模式", scene_label_zh(scene_id)),
            ("核心承诺 / Hook", hook_line or executive.get("conclusion", "").strip() or "核心承诺待补"),
            ("主证明镜头", proof_line or "主证明镜头待补"),
            ("制作交接", owner_line or executive.get("next_action", "").strip() or "制作交接待补"),
        ]
    if scene_id in CREATIVE_MATRIX_SCENE_IDS:
        matrix = sections.get("Variable Matrix") or {}
        handoff = sections.get("Execution Handoff") or sections.get("Production Handoff") or {}
        stage_line = first_table_cell(matrix, 0, 0) or ((matrix.get("table") or {}).get("title") or "")
        queue_line = first_table_cell(handoff, 0, 0) or first_table_cell(handoff, 1, 0)
        return [
            ("创意模式", scene_label_zh(scene_id)),
            ("当前主阶段", stage_line or executive.get("conclusion", "").strip() or "主阶段待补"),
            ("队列焦点", queue_line or "队列焦点待补"),
            ("下一步", executive.get("next_action", "").strip() or "下一步动作待补"),
        ]
    return [
        ("当前判断", executive.get("conclusion", "").strip() or "执行结论待补"),
        ("首要动作", executive.get("next_action", "").strip() or "下一步动作待补"),
        (
            "证据健康度",
            f"证据 {len(report.get('evidence', []) or [])} 条 / 资产 {len(report.get('assets', []) or [])} 条 / 来源 {len(report.get('sources', []) or [])} 条",
        ),
        (
            "执行优先级",
            "先看总览与章节概览，再从最强结构化章节进入动作落地。"
            if report.get("sections")
            else "当前报告仍需补章节内容。",
        ),
    ]


def scene_section_action_rows(report: dict, section: dict) -> list[tuple[str, str]]:
    scene_id = normalize_text(report.get("metadata", {}).get("scene", ""))
    heading = normalize_text(section.get("heading", ""))
    table = section.get("table") or {}
    if scene_id == "18" and heading == "Why They Matter":
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "周变化解释表"),
            ("为什么重要", "先判断谁在增强、谁在回落、谁只是事件噪音。"),
            ("本章焦点", "；".join(first_row[:2]) if len(first_row) >= 2 else "矩阵周变化待补"),
            ("交付去向", "用于继续追 / 借鉴 / 减少跟进 的周报分发"),
        ]
    if scene_id == "19" and heading == "High-Level Judgment":
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "高低表现主判断表"),
            ("为什么重要", "先把高表现、低表现和最佳发布时间窗拆开，再决定下轮排期。"),
            ("本章焦点", "；".join(first_row[:2]) if len(first_row) >= 2 else "高低表现对照待补"),
            ("交付去向", "用于多做 / 少做 / 停止 与下轮测试排期"),
        ]
    if scene_id in CREATIVE_BRIEF_SCENE_IDS and heading == "Production Handoff":
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "制作交接表"),
            ("为什么重要", "把锁定决策、责任人和阻塞风险一次性交给脚本、设计或渲染执行方。"),
            ("本章焦点", "；".join(cell for cell in first_row[:2] if normalize_text(cell)) or "制作交接项待补"),
            ("交付去向", "可继续生成 creative-production-handoff 运营包"),
        ]
    if scene_id in CREATIVE_BRIEF_SCENE_IDS and heading == "Message":
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "信息层改写表"),
            ("为什么重要", "先确认参考逻辑与适配版本是否分离，再进入镜头执行。"),
            ("本章焦点", "；".join(cell for cell in first_row[:2] if normalize_text(cell)) or "信息层待补"),
            ("交付去向", "用于锁定 hook、证明路径与 CTA"),
        ]
    if scene_id in CREATIVE_MATRIX_SCENE_IDS and heading == "Variable Matrix":
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "变量矩阵 / 阶段表"),
            ("为什么重要", "先看清阶段门槛、输入和产出，再决定本周该推进哪一段 pipeline。"),
            ("本章焦点", "；".join(cell for cell in first_row[:2] if normalize_text(cell)) or "阶段定义待补"),
            ("交付去向", "用于周度复制、测试或上新素材队列排期"),
        ]
    if scene_id in CREATIVE_MATRIX_SCENE_IDS and heading in {"Execution Handoff", "Production Handoff"}:
        first_row = (table.get("rows") or [[]])[0]
        return [
            ("先看什么", table.get("title") or "执行交接表"),
            ("为什么重要", "明确每个队列产物由谁拥有、何时可开工、卡在哪。"),
            ("本章焦点", "；".join(cell for cell in first_row[:2] if normalize_text(cell)) or "队列交接待补"),
            ("交付去向", "可继续生成 creative-production-handoff 运营包"),
        ]
    action_hint = table.get("title") or ("优先看结构化表格" if table.get("headers") else "优先看段落与要点")
    evidence_rows = len(section.get("evidence_refs", []) or [])
    evidence_hint = "已带证据引用，可直接回看来源。" if evidence_rows else "当前章节未带结构化证据引用。"
    return [
        ("先看什么", action_hint),
        ("为什么重要", section.get("instruction", "").strip() or "本章节用于承接当前场景的核心动作。"),
        ("证据状态", evidence_hint),
        ("执行判断", "章节可直接用于交接" if any([section.get("paragraphs"), section.get("bullets"), section.get("numbered"), table.get("headers")]) else "章节仍需补内容"),
    ]


def scene_summary_rows_xlsx(report: dict) -> list[tuple[str, str]]:
    scene_id = normalize_text(report.get("metadata", {}).get("scene", ""))
    rows = scene_cover_spotlight_rows(report)
    if scene_id == "18":
        return rows + [("阅读建议", "先看对象跟踪和关注理由，再看下一步动作分发。")]
    if scene_id == "19":
        return rows + [("阅读建议", "先看高低表现主判断，再看证据聚类与下轮测试计划。")]
    if scene_id in CREATIVE_BRIEF_SCENE_IDS:
        return rows + [("阅读建议", "先看 Message 与 Structure，再用 Production Handoff 进入制作执行。")]
    if scene_id in CREATIVE_MATRIX_SCENE_IDS:
        return rows + [("阅读建议", "先看 Variable Matrix，再用 Execution Handoff 推进周度队列。")]
    return rows


def repeat_doc_header_row(table) -> None:
    header = table.rows[0]._tr
    tr_pr = header.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn_shared("w:val"), "true")
    tr_pr.append(tbl_header)


def default_doc_table_widths(column_count: int) -> list[float]:
    width_map = {
        2: [1.6, 5.2],
        3: [1.4, 3.4, 2.0],
        4: [1.3, 2.4, 1.55, 1.55],
        5: [2.45, 0.8, 0.8, 0.8, 1.95],
        6: [1.75, 2.3, 0.6, 0.6, 0.6, 1.55],
    }
    if column_count in width_map:
        return width_map[column_count]
    usable_width = 6.8
    width = round(usable_width / max(column_count, 1), 2)
    return [width] * column_count


def set_doc_table_widths(table, column_widths: list[float] | None = None) -> None:
    widths = column_widths or default_doc_table_widths(len(table.columns))
    table.autofit = False
    for column_index, width in enumerate(widths[: len(table.columns)]):
        for cell in table.columns[column_index].cells:
            cell.width = Inches(width)


def add_doc_list(document: Document, title: str, values: list[str], style: str = "List Bullet") -> None:
    if not values:
        return
    document.add_heading(title, level=2)
    for item in values:
        document.add_paragraph(item, style=style)


def add_doc_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn_shared("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn_shared("xml:space"), "preserve")
    instr_text.text = instruction
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn_shared("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    placeholder = paragraph.add_run(" ")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn_shared("w:fldCharType"), "end")
    placeholder._r.append(fld_char_end)


def bookmark_name(text: str, index: int) -> str:
    normalized = re.sub(r"[^A-Za-z0-9_]+", "_", text).strip("_") or f"section_{index}"
    return f"sec_{index}_{normalized[:28]}"


def add_doc_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn_shared("w:id"), str(bookmark_id))
    start.set(qn_shared("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn_shared("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_doc_internal_link(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn_shared("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn_shared("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn_shared("w:val"), "single")
    run_pr.append(color)
    run_pr.append(underline)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(run_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_doc_toc(document: Document) -> None:
    heading = document.add_heading("目录", level=1)
    add_doc_bookmark(heading, "contents_anchor", 899)
    paragraph = document.add_paragraph()
    add_doc_field(paragraph, 'TOC \\o "1-2" \\h \\z \\u')
    note = document.add_paragraph()
    note_run = note.add_run("如果目录为空，请在 Word 中更新一次目录字段。")
    note_run.italic = True


def add_doc_navigation_links(document: Document, links: list[tuple[str, str]]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    for index, (label, anchor) in enumerate(links):
        add_doc_internal_link(paragraph, label, anchor)
        if index < len(links) - 1:
            paragraph.add_run(" | ")


def add_doc_header_footer(document: Document, report: dict) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.text = f"场景 {report['metadata'].get('scene', '')} | {localize_project_text(report['metadata'].get('project', ''))}"

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(report["metadata"].get("generated_at", "") or "生成时间待补")
    footer.add_run(" | 第 ")
    add_doc_field(footer, "PAGE")
    footer.add_run(" 页")


def add_doc_cover_page(document: Document, report: dict) -> None:
    metadata = report["metadata"]
    executive = report["executive_summary"]
    working_context = report["working_context"]
    theme = scene_doc_visual_theme(normalize_text(metadata.get("scene", "")))

    add_doc_kicker(document, theme["kicker"])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run(localized_report_title(metadata))
    title_run.bold = True
    title_run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.add_run(
        f"场景 {metadata.get('scene', '')} | {scene_label_zh(metadata.get('scene', ''))}".strip(" |")
    ).italic = True

    detail = document.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail.paragraph_format.space_after = Pt(18)
    detail.add_run(
        f"{localize_project_text(metadata.get('project', ''))} | {localize_status_text(metadata.get('status', '')) or '草稿'} | {metadata.get('generated_at', '') or '生成时间待补'}".strip(
            " |"
        )
    )

    add_doc_badge_row(
        document,
        [
            ("交付类型", normalize_text(metadata.get("deliverable_type", "")) or "TikTok 增长运营交付"),
            ("场景状态", localize_status_text(metadata.get("status", "")) or "草稿"),
            ("工作流标识", normalize_text(metadata.get("scene_slug", "")) or "待补"),
            ("章节数量", str(len(report.get("sections", []) or []))),
        ],
    )
    add_doc_metric_grid(
        document,
        [
            ("结构化章节", str(sum(1 for section in report.get("sections", []) if (section.get("table") or {}).get("headers")))),
            ("证据章节", str(sum(1 for section in report.get("sections", []) if section.get("evidence_refs")))),
            ("资产条目", str(len(report.get("assets", []) or []))),
            ("证据条目", str(len(report.get("evidence", []) or []))),
            ("备注条目", str(len(report.get("notes", []) or []))),
            ("执行模板", "已内置" if has_execution_template(report) else "未内置"),
        ],
    )
    add_doc_spotlight_table(
        document,
        "项目快照",
        "把这次交付当成平台项目卡而不是单次本地导出，先看关键动作、证据状态和执行优先级。",
        scene_cover_spotlight_rows(report),
        accent_fill=theme["accent"],
        widths=[1.45, 5.35],
    )

    overview_card = document.add_table(rows=2, cols=3)
    overview_card.cell(0, 0).text = "场景"
    overview_card.cell(0, 1).text = "状态"
    overview_card.cell(0, 2).text = "交付定位"
    overview_card.cell(1, 0).text = localized_scene_display(metadata)
    overview_card.cell(1, 1).text = localize_status_text(metadata.get("status", "")) or "草稿"
    overview_card.cell(1, 2).text = normalize_text(metadata.get("deliverable_type", "")) or "TikTok 增长运营交付"
    style_doc_card_table(overview_card, accent_fill=theme["card_accent"])
    for run in overview_card.cell(0, 0).paragraphs[0].runs + overview_card.cell(0, 1).paragraphs[0].runs + overview_card.cell(0, 2).paragraphs[0].runs:
        run.font.color.rgb = None
    set_doc_table_widths(overview_card, [2.0, 1.2, 3.6])

    cover_table = document.add_table(rows=4, cols=2)
    cover_rows = [
        ("场景", localized_scene_display(metadata)),
        ("交付物类型", metadata.get("deliverable_type", "")),
        ("状态", localize_status_text(metadata.get("status", "")) or "草稿"),
        ("工作流标识", metadata.get("scene_slug", "")),
    ]
    for index, (label, value) in enumerate(cover_rows):
        set_cell_text(cover_table.cell(index, 0), label)
        set_cell_text(cover_table.cell(index, 1), value)
    style_doc_table(cover_table, header_row=False, label_col=0)
    set_doc_table_widths(cover_table, [1.6, 5.2])

    add_doc_section_divider(document, "任务上下文", "本页汇总当前场景的输入背景、关键判断与执行入口。")
    context_lines = summary_preview_lines(working_context.get("summary", ""))
    if context_lines:
        for item in context_lines:
            paragraph = document.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
    else:
        context_body = document.add_paragraph("当前未提供任务上下文。")
        context_body.paragraph_format.space_after = Pt(12)

    add_doc_section_divider(document, "执行摘要", "优先展示本轮产物最值得看的结论、重要性与下一步动作。")
    summary_bits = [
        ("核心结论", executive.get("conclusion", "").strip()),
        ("为什么重要", executive.get("why_it_matters", "").strip()),
        ("下一步动作", executive.get("next_action", "").strip()),
    ]
    snapshot_table = document.add_table(rows=max(1, len([item for item in summary_bits if item[1]])), cols=2)
    snapshot_rows = [item for item in summary_bits if item[1]]
    if not snapshot_rows:
        snapshot_rows = [("核心结论", "执行摘要待补。")]
    for index, (label, value) in enumerate(snapshot_rows):
        set_cell_text(snapshot_table.cell(index, 0), label)
        set_cell_text(snapshot_table.cell(index, 1), value)
    style_doc_table(snapshot_table, header_row=False, label_col=0)
    set_doc_table_widths(snapshot_table, [1.35, 5.45])

    document.add_page_break()


def add_label_value(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(value or "N/A")


def render_table_docx(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    column_widths: list[float] | None = None,
) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    for column, header in enumerate(headers):
        set_cell_text(table.cell(0, column), header)
    for row in rows:
        cells = table.add_row().cells
        padded = row + [""] * max(0, len(headers) - len(row))
        for column, value in enumerate(padded[: len(headers)]):
            set_cell_text(cells[column], value)
    style_doc_table(table)
    repeat_doc_header_row(table)
    set_doc_table_widths(table, column_widths or infer_doc_table_widths(headers, rows))


def section_evidence_ref_rows(section: dict) -> list[list[str]]:
    rows = []
    for item in section.get("evidence_refs", []) or []:
        rows.append(
            [
                item.get("source_type", ""),
                item.get("source_id", ""),
                item.get("source_url", ""),
                item.get("time_range", ""),
                item.get("excerpt", ""),
                item.get("supports", ""),
            ]
        )
    return rows


def add_doc_image_caption(document: Document, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(caption)
    run.italic = True


def add_doc_scene_focus_card(document: Document, report: dict, section: dict, index: int) -> None:
    scene_id = normalize_text(report.get("metadata", {}).get("scene", ""))
    total_content_blocks = len(section.get("paragraphs", [])) + len(section.get("bullets", [])) + len(section.get("numbered", []))
    evidence_rows = len(section.get("evidence_refs", []) or [])
    table_rows = len(((section.get("table") or {}).get("rows")) or [])
    add_doc_section_divider(
        document,
        f"章节 {index:02d} · {section.get('display_heading', section.get('heading', ''))}",
        "以下卡片说明本章节的关注焦点、结构化程度与执行状态。",
    )
    card = document.add_table(rows=2, cols=5)
    card.cell(0, 0).text = "场景"
    card.cell(0, 1).text = "章节"
    card.cell(0, 2).text = "内容块"
    card.cell(0, 3).text = "结构化表格"
    card.cell(0, 4).text = "当前状态"
    card.cell(1, 0).text = scene_label_zh(scene_id)
    card.cell(1, 1).text = f"{index:02d} | {section.get('display_heading', section.get('heading', ''))}"
    card.cell(1, 2).text = str(total_content_blocks)
    card.cell(1, 3).text = section["table"]["title"] or ("已含结构化表格" if section["table"]["headers"] else "待补结构化表格")
    card.cell(1, 4).text = "章节可读" if any([section.get("paragraphs"), section.get("bullets"), section.get("numbered"), section["table"]["headers"]]) else "待补内容"
    style_doc_card_table(card, accent_fill="355C7D")
    set_doc_table_widths(card, [0.95, 2.35, 0.8, 1.7, 1.15])
    add_doc_badge_row(
        document,
        [
            ("段落块", str(len(section.get("paragraphs", [])))),
            ("要点块", str(len(section.get("bullets", [])))),
            ("步骤块", str(len(section.get("numbered", [])))),
            ("证据引用", str(evidence_rows)),
            ("表格行数", str(table_rows)),
        ],
    )
    add_doc_spotlight_table(
        document,
        "章节动作卡",
        "这张卡只回答三件事：先看什么、为什么重要、做完后交给谁。",
        scene_section_action_rows(report, section),
        accent_fill="355C7D",
        widths=[1.35, 5.45],
    )


def add_doc_section_overview(document: Document, report: dict, section_bookmarks: list[str]) -> None:
    heading = document.add_heading("章节概览", level=1)
    add_doc_bookmark(heading, "section_overview", 900)
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    intro.add_run("本页用于快速浏览本次报告的章节结构，并可直接跳转到对应章节。")
    add_doc_badge_row(
        document,
        [
            ("总章节数", str(len(report.get("sections", []) or []))),
            ("结构化章节", str(sum(1 for section in report.get("sections", []) if (section.get("table") or {}).get("headers")))),
            ("证据章节", str(sum(1 for section in report.get("sections", []) if section.get("evidence_refs")))),
        ],
    )
    table = document.add_table(rows=1, cols=5)
    headers = ["章节", "段落数", "要点数", "步骤数", "表格概况"]
    for column, header in enumerate(headers):
        set_cell_text(table.cell(0, column), header)
    for index, section in enumerate(report["sections"]):
        row = table.add_row().cells
        table_summary = section["table"]["title"] or ("已含结构化表格" if section["table"]["headers"] else "待补结构化表格")
        values = [
            section.get("display_heading", section["heading"]),
            str(len(section["paragraphs"])),
            str(len(section["bullets"])),
            str(len(section["numbered"])),
            table_summary,
        ]
        for column, value in enumerate(values):
            set_cell_text(row[column], value)
        paragraph = row[0].paragraphs[0]
        paragraph.clear()
        add_doc_internal_link(paragraph, section.get("display_heading", section["heading"]), section_bookmarks[index])
    style_doc_table(table)
    set_doc_table_widths(table, [2.45, 0.8, 0.8, 0.8, 1.95])


def add_doc_execution_template(document: Document, report: dict) -> None:
    if not has_execution_template(report):
        return
    execution_template = report["execution_template"]
    heading = document.add_heading("执行模板", level=1)
    heading.paragraph_format.keep_with_next = True

    recommended_request = execution_template.get("recommended_request", "")
    if recommended_request:
        add_label_value(document, "推荐请求词：", recommended_request)
    recommended_request_zh = execution_template.get("recommended_request_zh", "")
    if recommended_request_zh:
        add_label_value(document, "推荐中文请求词：", recommended_request_zh)

    runner_args = execution_template.get("recommended_runner_args", []) or []
    if runner_args:
        document.add_heading("运行参数", level=2)
        for item in runner_args:
            document.add_paragraph(item, style="List Bullet")

    variable_inputs = execution_template.get("variable_inputs", []) or []
    if variable_inputs:
        document.add_heading("可变输入", level=2)
        render_table_docx(
            document,
            ["变量", "含义", "示例", "必填"],
            [
                [item.get("name", ""), item.get("meaning", ""), item.get("example", ""), item.get("required", "")]
                for item in variable_inputs
            ],
            [1.25, 2.15, 2.6, 0.8],
        )

    prompt_scaffold = execution_template.get("codex_prompt_scaffold", []) or []
    if prompt_scaffold:
        document.add_heading("Codex 提示词骨架", level=2)
        for item in prompt_scaffold:
            document.add_paragraph(item, style="List Bullet")

    prompt_scaffold_zh = execution_template.get("codex_prompt_scaffold_zh", []) or []
    if prompt_scaffold_zh:
        document.add_heading("中文提示词骨架", level=2)
        for item in prompt_scaffold_zh:
            document.add_paragraph(item, style="List Bullet")

    workflow_steps = execution_template.get("workflow_steps", []) or []
    if workflow_steps:
        document.add_heading("执行步骤", level=2)
        for item in workflow_steps:
            document.add_paragraph(item, style="List Number")

    output_checklist = execution_template.get("output_checklist", []) or []
    if output_checklist:
        document.add_heading("交付检查清单", level=2)
        for item in output_checklist:
            document.add_paragraph(item, style="List Bullet")


def add_excel_table(ws, start_row: int, end_row: int, end_col: int, name: str) -> None:
    if end_row <= start_row or end_col < 1:
        return
    ref = f"A{start_row}:{get_column_letter(end_col)}{end_row}"
    table = Table(displayName=name, ref=ref)
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    ws.add_table(table)


def style_metric_card(cell, label: str, value: int) -> None:
    cell.value = f"{value}\n{label}"
    cell.font = Font(bold=True, color="FFFFFF", size=11)
    cell.fill = TITLE_FILL
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = THIN_BORDER


def style_status_metric_card(cell, label: str, value: int) -> None:
    if value <= 0:
        fill = SUCCESS_FILL
        font_color = "FFFFFF"
    elif value == 1:
        fill = WARNING_FILL
        font_color = "111111"
    else:
        fill = DANGER_FILL
        font_color = "FFFFFF"
    cell.value = f"{value}\n{label}"
    cell.font = Font(bold=True, color=font_color, size=11)
    cell.fill = fill
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell.border = THIN_BORDER


def write_docx(report: dict, output: Path) -> None:
    metadata = report["metadata"]
    working_context = report["working_context"]
    executive = report["executive_summary"]
    section_bookmarks = [bookmark_name(section["heading"], index) for index, section in enumerate(report["sections"], start=1)]

    document = Document()
    set_doc_font(document)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_doc_header_footer(document, report)
    add_doc_cover_page(document, report)
    title_heading = document.add_heading(localized_report_title(metadata), level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_table = document.add_table(rows=7, cols=2)
    meta_rows = [
        ("场景", localized_scene_display(metadata)),
        ("项目", localize_project_text(metadata.get("project", ""))),
        ("交付物类型", metadata.get("deliverable_type", "")),
        ("生成时间", metadata.get("generated_at", "")),
        ("状态", localize_status_text(metadata.get("status", ""))),
        ("场景说明文件", metadata.get("scenario_file", "")),
        ("场景标识", metadata.get("scene_slug", "")),
    ]
    for index, (label, value) in enumerate(meta_rows):
        set_cell_text(meta_table.cell(index, 0), label)
        set_cell_text(meta_table.cell(index, 1), value)
    style_doc_table(meta_table, header_row=False, label_col=0)
    set_doc_table_widths(meta_table, [1.6, 5.2])

    working_heading = document.add_heading("任务上下文", level=1)
    working_heading.paragraph_format.keep_with_next = True
    document.add_paragraph(working_context.get("summary", "") or "当前未提供任务上下文。")
    for key in [
        "inputs",
        "minimum_evidence",
        "ideal_evidence",
        "constraints",
        "requested_outputs",
        "ready_checklist",
    ]:
        label = CONTEXT_LABELS_ZH[key]
        add_doc_list(document, label, normalize_string_list(working_context.get(key)))

    executive_heading = document.add_heading("执行摘要", level=1)
    executive_heading.paragraph_format.keep_with_next = True
    add_label_value(document, "核心结论：", executive.get("conclusion", ""))
    add_label_value(document, "为什么重要：", executive.get("why_it_matters", ""))
    add_label_value(document, "下一步动作：", executive.get("next_action", ""))
    add_label_value(document, "置信度：", executive.get("confidence", ""))

    add_doc_toc(document)
    add_doc_section_overview(document, report, section_bookmarks)
    add_doc_navigation_links(
        document,
        [("返回目录", "contents_anchor"), ("跳到首个章节", section_bookmarks[0])] if section_bookmarks else [("返回目录", "contents_anchor")],
    )

    for label, key in [("操作检查清单", "operator_checklist"), ("常见失败模式", "common_failure_modes")]:
        add_doc_list(document, label, normalize_string_list(report["operator_guide"].get(key)))

    add_doc_execution_template(document, report)

    if report["evidence"]:
        evidence_heading = document.add_heading("证据总表", level=1)
        evidence_heading.paragraph_format.keep_with_next = True
        render_table_docx(
            document,
            ["标签", "详情", "来源"],
            [[item["label"], item["detail"], item["source"]] for item in report["evidence"]],
            [1.35, 3.95, 1.5],
        )

    if report["sections"]:
        document.add_page_break()

    for index, section in enumerate(report["sections"], start=1):
        if index > 1:
            document.add_page_break()
        heading = document.add_heading(section.get("display_heading", section["heading"]), level=1)
        heading.paragraph_format.keep_with_next = True
        add_doc_bookmark(heading, section_bookmarks[index - 1], 1000 + index)
        add_doc_navigation_links(document, [("返回目录", "contents_anchor"), ("返回章节概览", "section_overview")])
        add_doc_scene_focus_card(document, report, section, index)
        if section["instruction"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(section["instruction"])
            run.italic = True
        for paragraph_text in section["paragraphs"]:
            document.add_paragraph(paragraph_text)
        for item in section["bullets"]:
            document.add_paragraph(item, style="List Bullet")
        for item in section["numbered"]:
            document.add_paragraph(item, style="List Number")
        table = section["table"]
        if table["title"]:
            table_heading = document.add_heading(table["title"], level=2)
            table_heading.paragraph_format.keep_with_next = True
        section_layout = scene_section_layout(report, section)
        render_table_docx(document, table["headers"], table["rows"], section_layout.get("doc_widths"))
        evidence_ref_rows = section_evidence_ref_rows(section)
        if evidence_ref_rows:
            evidence_heading = document.add_heading("证据引用", level=2)
            evidence_heading.paragraph_format.keep_with_next = True
            render_table_docx(
                document,
                ["来源类型", "来源标识", "来源链接", "时间范围", "摘录", "支撑结论"],
                evidence_ref_rows,
            )
        if not any([section["paragraphs"], section["bullets"], section["numbered"], table["headers"], evidence_ref_rows]):
            document.add_paragraph("本章节待补充。")

    if report["assets"]:
        assets_heading = document.add_heading("资产清单", level=1)
        assets_heading.paragraph_format.keep_with_next = True
        render_table_docx(
            document,
            ["标签", "路径", "备注"],
            [[asset["label"], asset["path"], asset["note"]] for asset in report["assets"]],
            [1.3, 4.1, 1.4],
        )
        image_index = 1
        for asset in report["assets"]:
            path = Path(asset["path"]) if asset["path"] else None
            if path and path.exists() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp", ".gif"}:
                document.add_picture(str(path), width=Inches(5.6))
                add_doc_image_caption(document, f"Figure {image_index}. {asset['label'] or path.name}")
                image_index += 1

    if report["notes"]:
        notes_heading = document.add_heading("备注", level=1)
        notes_heading.paragraph_format.keep_with_next = True
        for item in report["notes"]:
            document.add_paragraph(item, style="List Bullet")

    if report["sources"]:
        sources_heading = document.add_heading("来源", level=1)
        sources_heading.paragraph_format.keep_with_next = True
        for item in report["sources"]:
            document.add_paragraph(item, style="List Bullet")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def style_title_row(ws, row: int, start_col: int, end_col: int, text: str, fill_hex: str = "") -> None:
    ws.merge_cells(start_row=row, start_column=start_col, end_row=row, end_column=end_col)
    cell = ws.cell(row=row, column=start_col)
    cell.value = text
    cell.font = Font(bold=True, color="FFFFFF", size=12)
    cell.fill = PatternFill(fill_type="solid", fgColor=fill_hex or "1F4E78")
    cell.alignment = Alignment(horizontal="center", vertical="center")
    cell.border = THIN_BORDER


def style_header_row(ws, row: int, headers: list[str]) -> None:
    for column, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=column)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = HEADER_FILL
        cell.alignment = WRAP
        cell.border = THIN_BORDER


def style_label_cell(cell) -> None:
    cell.font = Font(bold=True)
    cell.fill = LABEL_FILL
    cell.alignment = WRAP
    cell.border = THIN_BORDER


def style_value_cell(cell) -> None:
    cell.alignment = WRAP
    cell.border = THIN_BORDER


def apply_hyperlink(cell, target: str) -> None:
    if not target.strip():
        return
    cell.hyperlink = target
    cell.font = HYPERLINK_FONT


def add_sheet_back_link(
    ws,
    label: str = "返回章节概览",
    target: str = f"#'{get_sheet_title('section_overview')}'!A1",
) -> None:
    cell = ws.cell(row=2, column=1)
    cell.value = label
    cell.hyperlink = target
    cell.font = HYPERLINK_FONT


def finalize_sheet(ws, freeze_cell: str = "A4", filter_row: int | None = None) -> None:
    ws.freeze_panes = freeze_cell
    if filter_row is not None and ws.max_row >= filter_row and ws.max_column >= 1:
        ws.auto_filter.ref = f"A{filter_row}:{get_column_letter(ws.max_column)}{ws.max_row}"
    ws.sheet_view.zoomScale = 90


def safe_sheet_title(index: int, heading: str, used: set[str]) -> str:
    base_heading = localized_sheet_heading(heading)
    base = re.sub(r"[:\\/*?\[\]]+", "-", base_heading).strip() or f"Section {index}"
    prefix = f"{index:02d}-"
    budget = 31 - len(prefix)
    candidate = prefix + base[:budget]
    counter = 2
    while candidate in used:
        suffix = f"-{counter}"
        candidate = prefix + base[: max(1, budget - len(suffix))] + suffix
        counter += 1
    used.add(candidate)
    return candidate


def build_section_sheet_map(report: dict) -> list[str]:
    used_titles = set(RESERVED_SHEET_TITLES)
    return [
        safe_sheet_title(index, section["heading"], used_titles)
        for index, section in enumerate(report["sections"], start=1)
    ]


def asset_path_exists(value: str) -> bool:
    path_text = normalize_text(value or "")
    if not path_text:
        return False
    if not is_local_path_text(path_text):
        return True
    try:
        return Path(path_text).exists()
    except (OSError, ValueError):
        return False


def write_summary_sheet(workbook: Workbook, report: dict) -> None:
    ws = workbook.active
    ws.title = get_sheet_title("summary")
    theme = scene_doc_visual_theme(normalize_text(report.get("metadata", {}).get("scene", "")))
    style_title_row(ws, 1, 1, 6, localized_report_title(report["metadata"]), fill_hex=theme["title_fill"])
    metrics = [
        ("章节数", len(report["sections"])),
        ("证据条数", len(report["evidence"])),
        ("资产数", len(report["assets"])),
        ("备注数", len(report["notes"])),
        ("来源数", len(report["sources"])),
        ("开跑检查项", len(normalize_string_list(report["working_context"].get("ready_checklist")))),
    ]
    for column, (label, value) in enumerate(metrics, start=1):
        style_metric_card(ws.cell(row=3, column=column), label, value)
    ws.row_dimensions[3].height = 34
    quality_metrics = [
        ("空白章节", sum(1 for section in report["sections"] if not any([section["paragraphs"], section["bullets"], section["numbered"], section["table"]["headers"]]))),
        ("缺证据", 1 if not report["evidence"] else 0),
        ("缺资产", 1 if not report["assets"] else 0),
        ("失效资产路径", sum(1 for asset in report["assets"] if asset["path"] and not asset_path_exists(asset["path"]))),
        ("缺备注", 1 if not report["notes"] else 0),
        ("缺来源", 1 if not report["sources"] else 0),
    ]
    for column, (label, value) in enumerate(quality_metrics, start=1):
        style_status_metric_card(ws.cell(row=4, column=column), label, value)
    ws.row_dimensions[4].height = 34
    style_header_row(ws, 6, ["字段", "内容"])
    working_context_preview = " | ".join(summary_preview_lines(report["working_context"].get("summary", ""), limit=3))
    rows = [
        ("场景", localized_scene_display(report["metadata"])),
        ("项目", localize_project_text(report["metadata"].get("project", ""))),
        ("交付物类型", report["metadata"].get("deliverable_type", "")),
        ("生成时间", report["metadata"].get("generated_at", "")),
        ("状态", localize_status_text(report["metadata"].get("status", ""))),
        ("工作流标识", report["metadata"].get("scene_slug", "")),
        ("任务上下文", working_context_preview or report["working_context"].get("summary", "")),
        ("核心结论", report["executive_summary"].get("conclusion", "")),
        ("为什么重要", report["executive_summary"].get("why_it_matters", "")),
        ("下一步动作", report["executive_summary"].get("next_action", "")),
        ("置信度", report["executive_summary"].get("confidence", "")),
        (
            "证据健康度",
            f"证据 {len(report['evidence'])} 条 / 资产 {len(report['assets'])} 条 / 来源 {len(report['sources'])} 条",
        ),
        (
            "建议阅读路径",
            "先看章节概览，再看章节导航，最后进入结构化表格最重的章节。",
        ),
    ]
    rows.extend(scene_summary_rows_xlsx(report))
    row = 7
    for label, value in rows:
        ws.cell(row=row, column=1).value = label
        ws.cell(row=row, column=2).value = value
        style_label_cell(ws.cell(row=row, column=1))
        style_value_cell(ws.cell(row=row, column=2))
        set_sheet_row_height(ws, row, [label, value], chars_per_line=42, max_height=72)
        row += 1
    ws.cell(row=row, column=1).value = "章节概览"
    ws.cell(row=row, column=2).value = "打开章节概览页"
    style_label_cell(ws.cell(row=row, column=1))
    style_value_cell(ws.cell(row=row, column=2))
    apply_hyperlink(ws.cell(row=row, column=2), f"#'{get_sheet_title('section_overview')}'!A1")
    row += 1
    ws.cell(row=row, column=1).value = "章节导航"
    ws.cell(row=row, column=2).value = "打开章节到工作表的导航页"
    style_label_cell(ws.cell(row=row, column=1))
    style_value_cell(ws.cell(row=row, column=2))
    apply_hyperlink(ws.cell(row=row, column=2), f"#'{get_sheet_title('section_index')}'!A1")
    row += 1
    if has_execution_template(report):
        ws.cell(row=row, column=1).value = "执行模板"
        ws.cell(row=row, column=2).value = "打开可复用请求词、步骤和交付检查清单"
        style_label_cell(ws.cell(row=row, column=1))
        style_value_cell(ws.cell(row=row, column=2))
        apply_hyperlink(ws.cell(row=row, column=2), f"#'{get_sheet_title('execution_template')}'!A1")
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 78
    for column in range(3, 7):
        ws.column_dimensions[get_column_letter(column)].width = 15
    add_excel_table(ws, 6, row, 2, "SummaryTable")
    finalize_sheet(ws, freeze_cell="A7", filter_row=6)


def write_section_overview_sheet(workbook: Workbook, report: dict, section_sheet_map: list[str]) -> None:
    ws = workbook.create_sheet(get_sheet_title("section_overview"))
    style_title_row(ws, 1, 1, 8, "章节概览")
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    headers = ["章节", "章节状态", "证据状态", "填写说明", "段落数", "要点数", "步骤数", "表格标题"]
    style_header_row(ws, 3, headers)
    row = 4
    for index, section in enumerate(report["sections"], start=1):
        has_content = any([section["paragraphs"], section["bullets"], section["numbered"], section["table"]["headers"]])
        evidence_status = "带证据引用" if section.get("evidence_refs") else "无证据引用"
        values = [
            section.get("display_heading", section["heading"]),
            "章节可读" if has_content else "待补内容",
            evidence_status,
            section["instruction"],
            len(section["paragraphs"]),
            len(section["bullets"]),
            len(section["numbered"]),
            section["table"]["title"] or ("已含结构化表格" if section["table"]["headers"] else "待补结构化表格"),
        ]
        for column, value in enumerate(values, start=1):
            ws.cell(row=row, column=column).value = value
            style_value_cell(ws.cell(row=row, column=column))
        apply_hyperlink(ws.cell(row=row, column=1), f"#'{section_sheet_map[index - 1]}'!A1")
        row += 1
    widths = [28, 14, 14, 40, 12, 12, 12, 26]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    if row > 4:
        for column in ["E", "F", "G"]:
            ws.conditional_formatting.add(
                f"{column}4:{column}{row - 1}",
                ColorScaleRule(
                    start_type="min",
                    start_color="F2F7FC",
                    mid_type="percentile",
                    mid_value=50,
                    mid_color="9CC2E5",
                    end_type="max",
                    end_color="1F4E78",
                ),
            )
        add_excel_table(ws, 3, row - 1, 8, "SectionOverviewTable")
    finalize_sheet(ws, freeze_cell="A4", filter_row=3)


def write_section_index_sheet(workbook: Workbook, report: dict, section_sheet_map: list[str]) -> None:
    ws = workbook.create_sheet(get_sheet_title("section_index"))
    style_title_row(ws, 1, 1, 7, "章节导航")
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    headers = ["序号", "章节", "工作表", "填写说明", "段落数", "要点数", "步骤数"]
    style_header_row(ws, 3, headers)
    row = 4
    for index, (section, sheet_name) in enumerate(zip(report["sections"], section_sheet_map), start=1):
        values = [
            index,
            section.get("display_heading", section["heading"]),
            sheet_name,
            section["instruction"],
            len(section["paragraphs"]),
            len(section["bullets"]),
            len(section["numbered"]),
        ]
        for column, value in enumerate(values, start=1):
            ws.cell(row=row, column=column).value = value
            style_value_cell(ws.cell(row=row, column=column))
        apply_hyperlink(ws.cell(row=row, column=2), f"#'{sheet_name}'!A1")
        apply_hyperlink(ws.cell(row=row, column=3), f"#'{sheet_name}'!A1")
        row += 1
    widths = [8, 30, 26, 40, 12, 12, 12]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    if row > 4:
        for column in ["E", "F", "G"]:
            ws.conditional_formatting.add(
                f"{column}4:{column}{row - 1}",
                ColorScaleRule(
                    start_type="min",
                    start_color="F2F7FC",
                    mid_type="percentile",
                    mid_value=50,
                    mid_color="9CC2E5",
                    end_type="max",
                    end_color="1F4E78",
                ),
            )
        add_excel_table(ws, 3, row - 1, 7, "SectionIndexTable")
    finalize_sheet(ws, freeze_cell="A4", filter_row=3)


def write_operator_guide_sheet(workbook: Workbook, report: dict) -> None:
    ws = workbook.create_sheet(get_sheet_title("operator_guide"))
    style_title_row(ws, 1, 1, 2, "操作指引")
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    style_header_row(ws, 3, ["分类", "内容"])
    row = 4
    for label, key in [
        ("操作检查清单", "operator_checklist"),
        ("常见失败模式", "common_failure_modes"),
    ]:
        for item in normalize_string_list(report["operator_guide"].get(key)):
            ws.cell(row=row, column=1).value = label
            ws.cell(row=row, column=2).value = item
            style_label_cell(ws.cell(row=row, column=1))
            style_value_cell(ws.cell(row=row, column=2))
            row += 1
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 100
    if row > 4:
        add_excel_table(ws, 3, row - 1, 2, "OperatorGuideTable")
    finalize_sheet(ws, freeze_cell="A4", filter_row=3)


def write_context_lists_sheet(workbook: Workbook, report: dict) -> None:
    ws = workbook.create_sheet(get_sheet_title("context_lists"))
    style_title_row(ws, 1, 1, 2, "上下文清单")
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    style_header_row(ws, 3, ["分类", "内容"])
    row = 4
    for label, key in [
        ("输入材料", "inputs"),
        ("最低证据要求", "minimum_evidence"),
        ("理想证据补充", "ideal_evidence"),
        ("约束条件", "constraints"),
        ("目标交付", "requested_outputs"),
        ("开跑前检查", "ready_checklist"),
    ]:
        for item in normalize_string_list(report["working_context"].get(key)):
            ws.cell(row=row, column=1).value = label
            ws.cell(row=row, column=2).value = item
            style_label_cell(ws.cell(row=row, column=1))
            style_value_cell(ws.cell(row=row, column=2))
            row += 1
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 100
    if row > 4:
        add_excel_table(ws, 3, row - 1, 2, "ContextListsTable")
    finalize_sheet(ws, freeze_cell="A4", filter_row=3)


def write_execution_template_sheet(workbook: Workbook, report: dict) -> None:
    if not has_execution_template(report):
        return
    execution_template = report["execution_template"]
    ws = workbook.create_sheet(get_sheet_title("execution_template"))
    style_title_row(ws, 1, 1, 4, "执行模板")
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    row = 3

    for label, value in [
        ("推荐请求词", execution_template.get("recommended_request", "")),
        ("推荐中文请求词", execution_template.get("recommended_request_zh", "")),
    ]:
        if not value:
            continue
        ws.cell(row=row, column=1).value = label
        style_label_cell(ws.cell(row=row, column=1))
        ws.cell(row=row, column=2).value = value
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
        style_value_cell(ws.cell(row=row, column=2))
        set_sheet_row_height(ws, row, [label, value], chars_per_line=52)
        row += 1

    runner_args = execution_template.get("recommended_runner_args", []) or []
    if runner_args:
        style_title_row(ws, row, 1, 4, "运行参数")
        row += 1
        for item in runner_args:
            ws.cell(row=row, column=1).value = "运行参数"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = item
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [item], chars_per_line=52)
            row += 1

    variable_inputs = execution_template.get("variable_inputs", []) or []
    if variable_inputs:
        style_title_row(ws, row, 1, 4, "可变输入")
        row += 1
        header_row = row
        headers = ["变量", "含义", "示例", "必填"]
        style_header_row(ws, header_row, headers)
        row += 1
        variable_rows = []
        for item in variable_inputs:
            values = [item.get("name", ""), item.get("meaning", ""), item.get("example", ""), item.get("required", "")]
            variable_rows.append(values)
            for column, value in enumerate(values, start=1):
                ws.cell(row=row, column=column).value = value
                style_value_cell(ws.cell(row=row, column=column))
            set_sheet_row_height(ws, row, values, chars_per_line=28)
            row += 1
        add_excel_table(ws, header_row, row - 1, 4, "ExecutionTemplateVariables")
        widths = infer_sheet_column_widths(headers, variable_rows, min_width=12, max_width=34)
        for index, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(index)].width = width

    for title, items, numbered in [
        ("Codex 提示词骨架", execution_template.get("codex_prompt_scaffold", []) or [], False),
        ("中文提示词骨架", execution_template.get("codex_prompt_scaffold_zh", []) or [], False),
        ("执行步骤", execution_template.get("workflow_steps", []) or [], True),
        ("交付检查清单", execution_template.get("output_checklist", []) or [], False),
    ]:
        if not items:
            continue
        style_title_row(ws, row, 1, 4, title)
        row += 1
        for index, item in enumerate(items, start=1):
            ws.cell(row=row, column=1).value = f"步骤 {index}" if numbered else "条目"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = item
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [item], chars_per_line=52)
            row += 1

    for column, width in zip(["A", "B", "C", "D"], [22, 34, 34, 18]):
        ws.column_dimensions[column].width = max(ws.column_dimensions[column].width or 0, width)
    finalize_sheet(ws, freeze_cell="A3")


def write_section_sheets(workbook: Workbook, report: dict, section_sheet_map: list[str]) -> None:
    for index, (section, sheet_title) in enumerate(zip(report["sections"], section_sheet_map), start=1):
        ws = workbook.create_sheet(sheet_title)
        table = section["table"]
        section_layout = scene_section_layout(report, section)
        evidence_ref_headers = ["来源类型", "来源标识", "来源链接", "时间范围", "摘录", "支撑结论"]
        evidence_ref_rows = section_evidence_ref_rows(section)
        sheet_col_count = max(
            4,
            len(table["headers"]) if table["headers"] else 0,
            len(evidence_ref_headers) if evidence_ref_rows else 0,
        )
        style_title_row(ws, 1, 1, sheet_col_count, section.get("display_heading", section["heading"]))
        add_sheet_back_link(ws, target=f"#'{get_sheet_title('section_index')}'!A1")
        row = 3
        table_header_row: int | None = None
        table_end_row: int | None = None
        table_col_count = 0
        evidence_header_row: int | None = None
        evidence_end_row: int | None = None

        if section["instruction"]:
            ws.cell(row=row, column=1).value = "填写说明"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = section["instruction"]
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [section["instruction"]], chars_per_line=52)
            row += 2

        wrote_any = False
        for paragraph in section["paragraphs"]:
            ws.cell(row=row, column=1).value = "段落"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = paragraph
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [paragraph], chars_per_line=52)
            row += 1
            wrote_any = True

        for item in section["bullets"]:
            ws.cell(row=row, column=1).value = "要点"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = item
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [item], chars_per_line=52)
            row += 1
            wrote_any = True

        for order, item in enumerate(section["numbered"], start=1):
            ws.cell(row=row, column=1).value = f"步骤 {order}"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = item
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, [item], chars_per_line=52)
            row += 1
            wrote_any = True

        if table["headers"]:
            if table["title"]:
                ws.cell(row=row, column=1).value = "表格"
                style_label_cell(ws.cell(row=row, column=1))
                ws.cell(row=row, column=2).value = table["title"]
                ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
                style_value_cell(ws.cell(row=row, column=2))
                set_sheet_row_height(ws, row, [table["title"]], chars_per_line=52)
                row += 1
            table_header_row = row
            table_col_count = len(table["headers"])
            style_header_row(ws, row, table["headers"])
            row += 1
            for table_row in table["rows"]:
                padded = table_row + [""] * max(0, len(table["headers"]) - len(table_row))
                for column, value in enumerate(padded[: len(table["headers"])], start=1):
                    ws.cell(row=row, column=column).value = value
                    style_value_cell(ws.cell(row=row, column=column))
                set_sheet_row_height(
                    ws,
                    row,
                    padded[: len(table["headers"])],
                    chars_per_line=int(section_layout.get("xlsx_chars_per_line", 24)),
                )
                row += 1
            table_end_row = row - 1
            wrote_any = True

        if evidence_ref_rows:
            ws.cell(row=row, column=1).value = "证据引用"
            style_label_cell(ws.cell(row=row, column=1))
            ws.cell(row=row, column=2).value = "本章节所依赖的结构化来源证据。"
            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=2))
            set_sheet_row_height(ws, row, ["本章节所依赖的结构化来源证据。"], chars_per_line=52)
            row += 1
            evidence_header_row = row
            style_header_row(ws, row, evidence_ref_headers)
            row += 1
            for evidence_row in evidence_ref_rows:
                for column, value in enumerate(evidence_row, start=1):
                    ws.cell(row=row, column=column).value = value
                    style_value_cell(ws.cell(row=row, column=column))
                    if column == 3 and str(value).strip().startswith(("http://", "https://")):
                        apply_hyperlink(ws.cell(row=row, column=column), str(value).strip())
                set_sheet_row_height(ws, row, evidence_row, chars_per_line=28, max_height=90)
                row += 1
            evidence_end_row = row - 1
            wrote_any = True

        if not wrote_any:
            ws.cell(row=row, column=1).value = "本章节待补充。"
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=sheet_col_count)
            style_value_cell(ws.cell(row=row, column=1))

        if table["headers"]:
            table_widths = infer_sheet_column_widths(
                table["headers"],
                table["rows"],
                min_width=12,
                max_width=34,
                preferred_widths=list(section_layout.get("xlsx_preferred_widths", [])),
                per_column_max=list(section_layout.get("xlsx_per_column_max", [])),
            )
            ws.column_dimensions["A"].width = 18
            for column_index, width in enumerate(table_widths, start=1):
                letter = get_column_letter(column_index)
                ws.column_dimensions[letter].width = max(ws.column_dimensions[letter].width or 0, width)
        else:
            for column in range(1, sheet_col_count + 1):
                letter = get_column_letter(column)
                ws.column_dimensions[letter].width = 26 if column == 1 else 36
        if table_header_row is not None and table_end_row is not None and table_col_count > 0:
            add_excel_table(ws, table_header_row, table_end_row, table_col_count, f"SectionTable{index}")
        if evidence_header_row is not None and evidence_end_row is not None:
            add_excel_table(ws, evidence_header_row, evidence_end_row, len(evidence_ref_headers), f"SectionEvidenceRefs{index}")
        finalize_sheet(ws, freeze_cell="A3", filter_row=table_header_row)


def write_simple_list_sheet(workbook: Workbook, title: str, headers: list[str], rows: list[list[str]]) -> None:
    ws = workbook.create_sheet(title)
    layout = list_sheet_layout(title)
    chars_per_line = int(layout.get("chars_per_line", 26))
    preferred_widths = list(layout.get("preferred_widths", []))
    per_column_max = list(layout.get("per_column_max", []))
    style_title_row(ws, 1, 1, len(headers), title)
    add_sheet_back_link(ws, label="返回总览", target=f"#'{get_sheet_title('summary')}'!A1")
    style_header_row(ws, 3, headers)
    row = 4
    for values in rows:
        padded = values + [""] * max(0, len(headers) - len(values))
        for column, value in enumerate(padded[: len(headers)], start=1):
            ws.cell(row=row, column=column).value = value
            style_value_cell(ws.cell(row=row, column=column))
            text_value = str(value).strip()
            if text_value.startswith("http://") or text_value.startswith("https://"):
                apply_hyperlink(ws.cell(row=row, column=column), text_value)
            elif re.match(r"^[A-Za-z]:\\", text_value):
                apply_hyperlink(ws.cell(row=row, column=column), text_value)
        set_sheet_row_height(ws, row, padded[: len(headers)], chars_per_line=chars_per_line, max_height=84)
        row += 1
    widths = infer_sheet_column_widths(
        headers,
        rows,
        preferred_widths=preferred_widths,
        per_column_max=per_column_max,
    )
    for column_index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(column_index)].width = width
    if row > 4:
        safe_name = re.sub(r"[^A-Za-z0-9]+", "", title)
        if not safe_name:
            safe_name = f"Sheet{len(workbook.worksheets)}"
        add_excel_table(ws, 3, row - 1, len(headers), f"{safe_name}Table")
    finalize_sheet(ws, freeze_cell="A4", filter_row=3)


def write_report_artifact_sheets(workbook: Workbook, report: dict) -> None:
    board = report.get("collection_board")
    if isinstance(board, dict) and board.get("rows"):
        write_simple_list_sheet(
            workbook,
            "采集看板",
            board.get("headers") or [],
            board.get("rows") or [],
        )
    matrix = report.get("creation_matrix")
    if isinstance(matrix, dict):
        if matrix.get("matrix_rows"):
            write_simple_list_sheet(
                workbook,
                "创作就绪矩阵",
                matrix.get("matrix_headers") or [],
                matrix.get("matrix_rows") or [],
            )
        if matrix.get("pattern_rows"):
            write_simple_list_sheet(
                workbook,
                "共性规律收敛",
                matrix.get("pattern_headers") or [],
                matrix.get("pattern_rows") or [],
            )
    handoff = report.get("production_handoff")
    if isinstance(handoff, dict):
        tables = handoff.get("tables") or {}
        if tables.get("pacing_map"):
            write_simple_list_sheet(
                workbook,
                "Pacing Map",
                handoff.get("pacing_headers") or [],
                tables.get("pacing_map") or [],
            )
        if tables.get("subtitle_beats"):
            write_simple_list_sheet(
                workbook,
                "字幕节拍",
                handoff.get("subtitle_headers") or [],
                tables.get("subtitle_beats") or [],
            )
        if tables.get("proof_blocks"):
            write_simple_list_sheet(
                workbook,
                "证明块",
                handoff.get("proof_headers") or [],
                tables.get("proof_blocks") or [],
            )
        if tables.get("asset_requirements"):
            write_simple_list_sheet(
                workbook,
                "资产需求",
                handoff.get("asset_headers") or [],
                tables.get("asset_requirements") or [],
            )
        storyboard_rows = handoff.get("storyboard_rows") or []
        if storyboard_rows:
            write_simple_list_sheet(
                workbook,
                "Shot Handoff",
                ["shot_id", "时间", "阶段", "画面 / 动作", "字幕 / 口播", "generator 字段", "素材 / 执行需求"],
                storyboard_rows,
            )


def write_xlsx(report: dict, output: Path) -> None:
    workbook = Workbook()
    section_sheet_map = build_section_sheet_map(report)
    write_summary_sheet(workbook, report)
    write_section_overview_sheet(workbook, report, section_sheet_map)
    write_section_index_sheet(workbook, report, section_sheet_map)
    write_operator_guide_sheet(workbook, report)
    write_context_lists_sheet(workbook, report)
    write_execution_template_sheet(workbook, report)
    write_report_artifact_sheets(workbook, report)
    write_section_sheets(workbook, report, section_sheet_map)
    write_simple_list_sheet(
        workbook,
        get_sheet_title("evidence"),
        ["标签", "详情", "来源"],
        [[item["label"], item["detail"], item["source"]] for item in report["evidence"]],
    )
    write_simple_list_sheet(
        workbook,
        get_sheet_title("assets"),
        ["标签", "路径", "备注"],
        [[item["label"], item["path"], item["note"]] for item in report["assets"]],
    )
    write_simple_list_sheet(workbook, get_sheet_title("notes"), ["备注"], [[item] for item in report["notes"]])
    write_simple_list_sheet(workbook, get_sheet_title("sources"), ["来源"], [[item] for item in report["sources"]])
    output.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output)


def main() -> None:
    args = parse_args()
    report = resolve_payload(args)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    base_name = infer_base_name(report, args.base_name)
    formats = [item.strip().lower() for item in args.formats.split(",") if item.strip()]

    written: dict[str, str] = {}
    if "md" in formats:
        path = build_output_path(output_dir, base_name, ".md")
        write_utf8_text(path, render_localized_markdown(report))
        written["md"] = str(path)
    if "docx" in formats:
        path = build_output_path(output_dir, base_name, ".docx")
        write_docx(report, path)
        written["docx"] = str(path)
    if "xlsx" in formats:
        path = build_output_path(output_dir, base_name, ".xlsx")
        write_xlsx(report, path)
        written["xlsx"] = str(path)

    print(json.dumps(written, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
